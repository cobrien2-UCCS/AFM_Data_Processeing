from __future__ import annotations

import argparse
import csv
from collections import defaultdict
from datetime import datetime
from pathlib import Path
import statistics as stats
from statistics import mean, median

from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt

SOURCE_MODULUS_UNIT = "kPa"
DISPLAY_MODULUS_UNIT = "GPa-equivalent"
DISPLAY_MODULUS_SCALE = 1e-6


def _read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def _f(value: str | None) -> float:
    if value in (None, "", "nan", "NaN"):
        return float("nan")
    return float(value)


def _i(value: str | None) -> int:
    if value in (None, ""):
        return 0
    return int(float(value))


def _add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]]) -> None:
    p = doc.add_paragraph(title)
    if p.runs:
        p.runs[0].bold = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def _add_figure(doc: Document, title: str, path: Path, width: float = 6.0) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.add_paragraph(title)
        doc.add_paragraph(f"Source: {path}")
    else:
        doc.add_paragraph(f"[MISSING FIGURE] {title}")
        doc.add_paragraph(f"Expected source: {path}")


def _add_heading_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True


def _plot_boxplot(out_path: Path, title: str, labels: list[str], series: list[list[float]], ylabel: str) -> None:
    if not series or not any(s for s in series):
        return
    fig, ax = plt.subplots(figsize=(10.5, 5.5))
    ax.boxplot(series, tick_labels=labels, showfliers=False)
    ax.set_title(title)
    ax.set_ylabel(ylabel)
    ax.tick_params(axis="x", rotation=35)
    fig.tight_layout()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, dpi=300)
    plt.close(fig)


def _plot_bar_with_error(
    out_path: Path,
    title: str,
    labels: list[str],
    means: list[float],
    stds: list[float],
    ylabel: str,
) -> None:
    if not means:
        return
    fig, ax = plt.subplots(figsize=(10.5, 5.5))
    x = range(len(labels))
    ax.bar(x, means, yerr=stds, capsize=4, color="#4C78A8", edgecolor="black")
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels, rotation=30, ha="right")
    ax.set_title(title)
    ax.set_ylabel(ylabel)
    fig.tight_layout()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, dpi=300)
    plt.close(fig)


def _plot_combined_baseline_heatmap(out_path: Path, paired_rows: list[dict[str, str]]) -> None:
    rows = [r for r in paired_rows if r.get("method") == "gwy_stats"]
    if not rows:
        return
    max_row = max(_i(r.get("row_idx")) for r in rows)
    max_col = max(_i(r.get("col_idx")) for r in rows)
    import numpy as np

    grid = np.full((max_row + 1, max_col + 1), np.nan)
    for row in rows:
        rr = _i(row.get("row_idx"))
        cc = _i(row.get("col_idx"))
        fwd = _f(row.get("avg_forward"))
        bwd = _f(row.get("avg_backward"))
        if fwd == fwd and bwd == bwd:
            grid[rr, cc] = ((fwd + bwd) / 2.0) * DISPLAY_MODULUS_SCALE

    cmap = plt.cm.get_cmap("viridis").copy()
    cmap.set_bad(color="lightgray")
    fig, ax = plt.subplots(figsize=(6.2, 5.2))
    im = ax.imshow(grid, origin="upper", cmap=cmap)
    ax.set_title("Combined Forward/Backward Baseline Modulus\nMean of paired baseline fields")
    ax.set_xlabel("Grid column")
    ax.set_ylabel("Grid row")
    cbar = fig.colorbar(im, ax=ax)
    cbar.set_label(f"Modulus ({DISPLAY_MODULUS_UNIT})")
    fig.tight_layout()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, dpi=300)
    plt.close(fig)


def _aggregate_compare_rows(rows: list[dict[str, str]]) -> dict[str, dict[str, float]]:
    grouped: dict[str, dict[str, list[float]]] = defaultdict(lambda: defaultdict(list))
    for row in rows:
        method = row["method"]
        avg = _f(row.get("avg_value"))
        avg_baseline = _f(row.get("avg_baseline"))
        std = _f(row.get("std_value"))
        std_baseline = _f(row.get("std_baseline"))
        n_valid = _f(row.get("n_valid"))
        n_valid_baseline = _f(row.get("n_valid_baseline"))
        delta_avg = _f(row.get("delta_avg"))
        delta_std = _f(row.get("delta_std"))
        delta_n = _f(row.get("delta_n_valid"))
        if avg_baseline == avg_baseline and avg_baseline != 0.0 and avg == avg:
            grouped[method]["ratio_avg"].append(avg / avg_baseline)
        if std_baseline == std_baseline and std_baseline != 0.0 and std == std:
            grouped[method]["ratio_std"].append(std / std_baseline)
        if n_valid_baseline == n_valid_baseline and n_valid_baseline != 0.0 and n_valid == n_valid:
            grouped[method]["ratio_n"].append(n_valid / n_valid_baseline)
        if delta_avg == delta_avg:
            grouped[method]["delta_avg"].append(delta_avg)
        if delta_std == delta_std:
            grouped[method]["delta_std"].append(delta_std)
        if delta_n == delta_n:
            grouped[method]["delta_n"].append(delta_n)

    out: dict[str, dict[str, float]] = {}
    for method, metrics in grouped.items():
        out[method] = {
            "n_scans": len(metrics["delta_avg"]),
            "mean_ratio_avg": mean(metrics["ratio_avg"]) if metrics["ratio_avg"] else float("nan"),
            "median_ratio_avg": median(metrics["ratio_avg"]) if metrics["ratio_avg"] else float("nan"),
            "mean_delta_avg": mean(metrics["delta_avg"]) if metrics["delta_avg"] else float("nan"),
            "mean_ratio_std": mean(metrics["ratio_std"]) if metrics["ratio_std"] else float("nan"),
            "mean_delta_std": mean(metrics["delta_std"]) if metrics["delta_std"] else float("nan"),
            "mean_delta_n": mean(metrics["delta_n"]) if metrics["delta_n"] else float("nan"),
            "min_delta_n": min(metrics["delta_n"]) if metrics["delta_n"] else float("nan"),
            "max_delta_n": max(metrics["delta_n"]) if metrics["delta_n"] else float("nan"),
        }
    return out


def _absolute_method_summary(rows: list[dict[str, str]]) -> dict[str, dict[str, float]]:
    grouped: dict[str, dict[str, list[float]]] = defaultdict(lambda: defaultdict(list))
    for row in rows:
        method = row["method"]
        avg = _f(row.get("avg_value"))
        std = _f(row.get("std_value"))
        n_valid = _f(row.get("n_valid"))
        if avg == avg:
            grouped[method]["avg_value"].append(avg)
        if std == std:
            grouped[method]["std_value"].append(std)
        if n_valid == n_valid:
            grouped[method]["n_valid"].append(n_valid)
    out: dict[str, dict[str, float]] = {}
    for method, metrics in grouped.items():
        avgs = metrics.get("avg_value", [])
        stds = metrics.get("std_value", [])
        nvals = metrics.get("n_valid", [])
        out[method] = {
            "n_scans": len(avgs),
            "mean_avg_value": mean(avgs) if avgs else float("nan"),
            "std_avg_value": stats.pstdev(avgs) if len(avgs) > 1 else 0.0,
            "se_avg_value": (stats.pstdev(avgs) / (len(avgs) ** 0.5)) if len(avgs) > 1 else 0.0,
            "mean_std_value": mean(stds) if stds else float("nan"),
            "std_std_value": stats.pstdev(stds) if len(stds) > 1 else 0.0,
            "mean_n_valid": mean(nvals) if nvals else float("nan"),
        }
    return out


def _pixel_loss_summary(rows: list[dict[str, str]]) -> dict[str, dict[str, float]]:
    grouped: dict[str, dict[str, list[float]]] = defaultdict(lambda: defaultdict(list))
    for row in rows:
        method = row["method"]
        delta_n = _f(row.get("delta_n_valid"))
        n_valid_baseline = _f(row.get("n_valid_baseline"))
        if delta_n != delta_n:
            continue
        pixels_lost = max(0.0, -delta_n)
        grouped[method]["pixels_lost"].append(pixels_lost)
        if n_valid_baseline == n_valid_baseline and n_valid_baseline > 0:
            grouped[method]["pct_lost"].append(100.0 * pixels_lost / n_valid_baseline)
    out: dict[str, dict[str, float]] = {}
    for method, metrics in grouped.items():
        vals = metrics.get("pixels_lost", [])
        pct_vals = metrics.get("pct_lost", [])
        out[method] = {
            "n_scans": len(vals),
            "mean_pixels_lost": mean(vals) if vals else float("nan"),
            "std_pixels_lost": stats.pstdev(vals) if len(vals) > 1 else 0.0,
            "se_pixels_lost": (stats.pstdev(vals) / (len(vals) ** 0.5)) if len(vals) > 1 else 0.0,
            "max_pixels_lost": max(vals) if vals else float("nan"),
            "scans_with_loss": sum(1 for v in vals if v > 0),
            "mean_pct_lost": mean(pct_vals) if pct_vals else float("nan"),
            "std_pct_lost": stats.pstdev(pct_vals) if len(pct_vals) > 1 else 0.0,
            "se_pct_lost": (stats.pstdev(pct_vals) / (len(pct_vals) ** 0.5)) if len(pct_vals) > 1 else 0.0,
            "max_pct_lost": max(pct_vals) if pct_vals else float("nan"),
            "scans_gt_5pct": sum(1 for v in pct_vals if v > 5.0),
            "scans_gt_10pct": sum(1 for v in pct_vals if v > 10.0),
            "scans_gt_25pct": sum(1 for v in pct_vals if v > 25.0),
        }
    return out


def _load_baseline_inventory(path: Path) -> tuple[int, str]:
    rows = _read_csv(path)
    units = rows[0].get("units", "") if rows else ""
    return len(rows), units


def _fmt_ratio(value: float) -> str:
    return "" if value != value else f"{value:.3f}"


def _fmt_num(value: float) -> str:
    return "" if value != value else f"{value:.3e}"


def _fmt_modulus_display(value: float) -> str:
    return "" if value != value else f"{value * DISPLAY_MODULUS_SCALE:.3f}"


def build_report(
    forward_compare_dir: Path,
    backward_compare_dir: Path,
    paired_dir: Path,
    forward_baseline_summary: Path,
    backward_baseline_summary: Path,
    output_docx: Path,
) -> None:
    forward_long = _read_csv(forward_compare_dir / "comparison_long.csv")
    backward_long = _read_csv(backward_compare_dir / "comparison_long.csv")
    paired_rows = _read_csv(paired_dir / "fwd_bwd_summary.csv")
    paired_long = _read_csv(paired_dir / "fwd_bwd_long.csv")
    forward_baseline_rows = _read_csv(forward_baseline_summary)
    backward_baseline_rows = _read_csv(backward_baseline_summary)

    forward_agg = _aggregate_compare_rows(forward_long)
    backward_agg = _aggregate_compare_rows(backward_long)
    forward_abs = _absolute_method_summary(forward_long)
    backward_abs = _absolute_method_summary(backward_long)
    forward_loss = _pixel_loss_summary(forward_long)
    backward_loss = _pixel_loss_summary(backward_long)
    n_forward, units_forward = _load_baseline_inventory(forward_baseline_summary)
    n_backward, units_backward = _load_baseline_inventory(backward_baseline_summary)
    output_dir = output_docx.parent
    asset_dir = output_dir / "modulus_baseline_assets"
    asset_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("Modulus Baseline Validation Report", level=0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(
        "This report summarizes the baseline PEGDA modulus method-comparison results used to support Chapter 6, "
        "Section 6.1 (Baseline PEGDA Validation). Its purpose is to document forward/backward agreement, route "
        "consistency, and the practical basis for restricting the Stage 1 topography workflow to forward scans."
    )
    doc.add_paragraph(
        f"Unit note: the source `summary.csv` files for this study record modulus in {SOURCE_MODULUS_UNIT}. "
        f"To improve readability, modulus-derived values in the tables and plots below are displayed in "
        f"{DISPLAY_MODULUS_UNIT} by dividing the source-reported {SOURCE_MODULUS_UNIT} values by 1,000,000. "
        f"The baseline inventory table still reports the source unit exactly as provided by the pipeline. "
        f"For the current TIFF path, that source unit should be treated as the active workflow label, not yet as independently confirmed instrument metadata."
    )
    doc.add_paragraph(
        "Verification note: a direct one-file runner check on the current modulus TIFFs showed that pygwy did not detect "
        "embedded z-units for that file (`get_si_unit_z()` returned empty), and the runner therefore populated the output "
        "unit from the workflow default/fallback path. The current `kPa` labeling should therefore be treated as the "
        "active workflow assumption for these TIFFs, not yet as independently validated source metadata truth."
    )
    doc.add_paragraph(
        "Additional caution: that same one-file verification run also exported a negative modulus value. The broader "
        "modulus method-comparison CSV sets used in this report remained non-negative, so the comparative route-validation "
        "results are still usable. However, the absolute modulus value path should be treated as provisional pending "
        "targeted re-validation of the TIFF unit/processing chain."
    )

    _add_heading_paragraph(doc, "Workflow and method definitions")
    doc.add_paragraph(
        "The modulus baseline workflow compares several processing/statistics routes applied to the same PEGDA modulus "
        "scan set. The baseline route is the Gwyddion-statistics path (`modulus_gwy_stats` / `gwy_stats`), which uses "
        "Gwyddion preprocessing and Gwyddion-derived summary statistics. Alternative routes keep the same overall scan set "
        "but change how the post-preprocessing statistics or raw-value filtering are performed."
    )
    _add_table(
        doc,
        "Table M0 - Method definitions used in the modulus baseline validation",
        ["Method", "Definition", "Purpose in this report"],
        [
            ["modulus_gwy_stats (baseline)", "Gwyddion preprocessing with Gwyddion-derived statistics.", "Reference route for all modulus comparisons."],
            ["gwy_ops_py_stats", "Gwyddion preprocessing with Python-derived statistics on the processed field.", "Checks whether the stats engine changes the modulus summary materially."],
            ["raw_minmax", "Raw-value export with Python min/max validity filtering.", "Checks a simple deterministic raw-data filtering route."],
            ["raw_chauvenet", "Raw-value export with Chauvenet outlier filtering.", "Checks sensitivity to probabilistic outlier rejection."],
            ["raw_three_sigma", "Raw-value export with 3-sigma outlier filtering.", "Checks sensitivity to sigma-based outlier rejection."],
            ["Forward/backward paired summary", "Row/column-matched comparison between forward and backward scans.", "Checks directional consistency and supports forward-only selection in the topo Stage 1 workflow."],
        ],
    )
    _add_heading_paragraph(doc, "Calculation definitions used in the modulus comparison study")
    doc.add_paragraph(
        "For each scan, the comparison script reads the reported modulus summary from each method and compares it to the "
        "baseline Gwyddion-statistics route. The per-scan quantities used in this report are:"
    )
    doc.add_paragraph("avg_value: mean modulus reported by the comparison method for a scan.")
    doc.add_paragraph("std_value: standard deviation of modulus reported by the comparison method for a scan.")
    doc.add_paragraph("n_valid: number of valid pixels retained for that scan.")
    doc.add_paragraph("avg_baseline: mean modulus from the baseline route for the same scan.")
    doc.add_paragraph("std_baseline: standard deviation from the baseline route for the same scan.")
    doc.add_paragraph("n_valid_baseline: valid-pixel count from the baseline route for the same scan.")
    doc.add_paragraph("delta_avg = avg_value - avg_baseline")
    doc.add_paragraph("delta_std = std_value - std_baseline")
    doc.add_paragraph("delta_n_valid = n_valid - n_valid_baseline")
    doc.add_paragraph(
        "For the paired forward/backward comparison, scans are matched by grid row/column index and the paired quantities are:"
    )
    doc.add_paragraph("delta_avg(F/B) = avg_backward - avg_forward")
    doc.add_paragraph("ratio_avg(F/B) = avg_backward / avg_forward")
    doc.add_paragraph(
        "A ratio near 1.0 indicates close directional agreement. A delta near 0 indicates close agreement with the "
        "baseline route or, in the paired case, close agreement between scan directions."
    )
    doc.add_paragraph(
        "Important interpretation note: the absolute-modulus summary tables report the spread of scan-level mean modulus "
        "values across the whole dataset for each method. That is a between-scan variability measure. It is different "
        "from the within-scan spatial standard deviation shown in individual scan summaries or from the visually smoother "
        "mean heatmaps."
    )
    doc.add_paragraph(
        "To keep the bar plots readable, the bar-plot error bars in this report use standard error rather than full "
        "standard deviation. The full standard deviation values remain listed in the summary tables."
    )
    doc.add_paragraph(
        "In the grouped bar plots, each bar summarizes a method-level modulus comparison against the baseline route. In the "
        "heatmaps, the baseline heatmap shows the spatial distribution of the baseline modulus summaries, while the "
        "delta-versus-baseline heatmaps show how a comparison route shifts the spatial pattern relative to that same baseline."
    )
    doc.add_paragraph(
        "Outlined cells indicate scan positions where the comparison method retained fewer valid pixels than the baseline "
        "route. The outline is drawn in dark gray so it can be distinguished from the red-blue diverging color scale used "
        "for modulus differences."
    )
    doc.add_paragraph(
        "A visually quiet or nearly blank delta heatmap does not necessarily indicate missing data. In several cases, "
        "particularly for the raw min/max route, the comparison values remain very close to the baseline or to the "
        "Gwyddion-preprocess/Python-stats route, so the spatial delta field compresses toward zero."
    )

    _add_table(
        doc,
        "Table M1 - Baseline dataset context",
        ["Dataset", "Baseline method", "Scans", "Units", "Compare directory"],
        [
            ["Forward modulus baseline", "modulus_gwy_stats", str(n_forward), units_forward, str(forward_compare_dir)],
            ["Backward modulus baseline", "modulus_gwy_stats", str(n_backward), units_backward, str(backward_compare_dir)],
            ["Paired forward/backward summary", "paired row/col matches", str(len(paired_rows)), "", str(paired_dir)],
        ],
    )
    _add_table(
        doc,
        "Table M1a - Baseline file/source inventory",
        ["Direction", "Scan count", "Example source file", "Baseline summary path"],
        [
            [
                "Forward",
                str(n_forward),
                (forward_baseline_rows[0]["source_file"] if forward_baseline_rows else ""),
                str(forward_baseline_summary),
            ],
            [
                "Backward",
                str(n_backward),
                (backward_baseline_rows[0]["source_file"] if backward_baseline_rows else ""),
                str(backward_baseline_summary),
            ],
        ],
    )
    _add_table(
        doc,
        "Table M1b - Unit provenance interpretation used in this report",
        ["Output family", "Current unit behavior", "Interpretation"],
        [
            ["Modulus source summaries", "Pipeline outputs currently show kPa", "Treat as workflow-level fallback/normalization unless verified against instrument/export metadata; absolute values remain provisional."],
            ["Particle count outputs", "Count-based; no physical z-unit", "Unitless counts of retained or isolated particles."],
            ["Particle/grain diameter outputs", "nm encoded in metric/column names", "Derived geometric metrics, not inherited z-units from the TIFF field."],
        ],
    )

    method_order = ["gwy_ops_py_stats", "raw_minmax", "raw_chauvenet", "raw_three_sigma"]
    method_labels = {
        "gwy_ops_py_stats": "Gwyddion preprocess + Python stats",
        "raw_minmax": "Raw export + min/max filter",
        "raw_chauvenet": "Raw export + Chauvenet",
        "raw_three_sigma": "Raw export + 3-sigma",
        "gwy_stats": "Forward/backward paired baseline",
    }

    _add_table(
        doc,
        "Table M2 - Forward method comparison summary",
        [
            "Method",
            "Scans",
            "Mean ratio vs baseline",
            "Median ratio vs baseline",
            f"Mean delta avg ({DISPLAY_MODULUS_UNIT})",
            "Mean delta n_valid",
        ],
        [
            [
                method_labels.get(m, m),
                str(int(forward_agg[m]["n_scans"])),
                _fmt_ratio(forward_agg[m]["mean_ratio_avg"]),
                _fmt_ratio(forward_agg[m]["median_ratio_avg"]),
                _fmt_modulus_display(forward_agg[m]["mean_delta_avg"]),
                _fmt_num(forward_agg[m]["mean_delta_n"]),
            ]
            for m in method_order
            if m in forward_agg
        ],
    )

    _add_table(
        doc,
        "Table M3 - Backward method comparison summary",
        [
            "Method",
            "Scans",
            "Mean ratio vs baseline",
            "Median ratio vs baseline",
            f"Mean delta avg ({DISPLAY_MODULUS_UNIT})",
            "Mean delta n_valid",
        ],
        [
            [
                method_labels.get(m, m),
                str(int(backward_agg[m]["n_scans"])),
                _fmt_ratio(backward_agg[m]["mean_ratio_avg"]),
                _fmt_ratio(backward_agg[m]["median_ratio_avg"]),
                _fmt_modulus_display(backward_agg[m]["mean_delta_avg"]),
                _fmt_num(backward_agg[m]["mean_delta_n"]),
            ]
            for m in method_order
            if m in backward_agg
        ],
    )
    _add_table(
        doc,
        "Table M3a - Absolute modulus summary by method",
        [
            "Direction",
            "Method",
            "Scans",
            f"Mean modulus ({DISPLAY_MODULUS_UNIT})",
            f"Std of scan means across dataset ({DISPLAY_MODULUS_UNIT})",
            f"Mean per-scan std ({DISPLAY_MODULUS_UNIT})",
            "Mean n_valid",
        ],
        [
            [
                "Forward",
                method_labels.get(m, m),
                str(int(forward_abs[m]["n_scans"])),
                _fmt_modulus_display(forward_abs[m]["mean_avg_value"]),
                _fmt_modulus_display(forward_abs[m]["std_avg_value"]),
                _fmt_modulus_display(forward_abs[m]["mean_std_value"]),
                _fmt_num(forward_abs[m]["mean_n_valid"]),
            ]
            for m in method_order
            if m in forward_abs
        ] + [
            [
                "Backward",
                method_labels.get(m, m),
                str(int(backward_abs[m]["n_scans"])),
                _fmt_modulus_display(backward_abs[m]["mean_avg_value"]),
                _fmt_modulus_display(backward_abs[m]["std_avg_value"]),
                _fmt_modulus_display(backward_abs[m]["mean_std_value"]),
                _fmt_num(backward_abs[m]["mean_n_valid"]),
            ]
            for m in method_order
            if m in backward_abs
        ],
    )
    _add_table(
        doc,
        "Table M3b - Pixels lost relative to baseline by method",
        ["Direction", "Method", "Mean pixels lost", "Std pixels lost", "Max pixels lost", "Scans with any loss"],
        [
            [
                "Forward",
                method_labels.get(m, m),
                _fmt_num(forward_loss[m]["mean_pixels_lost"]),
                _fmt_num(forward_loss[m]["std_pixels_lost"]),
                _fmt_num(forward_loss[m]["max_pixels_lost"]),
                str(int(forward_loss[m]["scans_with_loss"])),
            ]
            for m in method_order
            if m in forward_loss
        ] + [
            [
                "Backward",
                method_labels.get(m, m),
                _fmt_num(backward_loss[m]["mean_pixels_lost"]),
                _fmt_num(backward_loss[m]["std_pixels_lost"]),
                _fmt_num(backward_loss[m]["max_pixels_lost"]),
                str(int(backward_loss[m]["scans_with_loss"])),
            ]
            for m in method_order
            if m in backward_loss
        ],
    )
    _add_table(
        doc,
        "Table M3c - Percent pixels lost relative to baseline by method",
        ["Direction", "Method", "Mean % lost", "Std % lost", "Max % lost", "Scans >5%", "Scans >10%", "Scans >25%"],
        [
            [
                "Forward",
                method_labels.get(m, m),
                _fmt_ratio(forward_loss[m]["mean_pct_lost"]),
                _fmt_ratio(forward_loss[m]["std_pct_lost"]),
                _fmt_ratio(forward_loss[m]["max_pct_lost"]),
                str(int(forward_loss[m]["scans_gt_5pct"])),
                str(int(forward_loss[m]["scans_gt_10pct"])),
                str(int(forward_loss[m]["scans_gt_25pct"])),
            ]
            for m in method_order
            if m in forward_loss
        ] + [
            [
                "Backward",
                method_labels.get(m, m),
                _fmt_ratio(backward_loss[m]["mean_pct_lost"]),
                _fmt_ratio(backward_loss[m]["std_pct_lost"]),
                _fmt_ratio(backward_loss[m]["max_pct_lost"]),
                str(int(backward_loss[m]["scans_gt_5pct"])),
                str(int(backward_loss[m]["scans_gt_10pct"])),
                str(int(backward_loss[m]["scans_gt_25pct"])),
            ]
            for m in method_order
            if m in backward_loss
        ],
    )
    _add_table(
        doc,
        "Table M3b - Pixels lost relative to baseline by method",
        ["Direction", "Method", "Mean pixels lost", "Std pixels lost", "Max pixels lost", "Scans with any loss"],
        [
            [
                "Forward",
                method_labels.get(m, m),
                _fmt_num(forward_loss[m]["mean_pixels_lost"]),
                _fmt_num(forward_loss[m]["std_pixels_lost"]),
                _fmt_num(forward_loss[m]["max_pixels_lost"]),
                str(int(forward_loss[m]["scans_with_loss"])),
            ]
            for m in method_order
            if m in forward_loss
        ] + [
            [
                "Backward",
                method_labels.get(m, m),
                _fmt_num(backward_loss[m]["mean_pixels_lost"]),
                _fmt_num(backward_loss[m]["std_pixels_lost"]),
                _fmt_num(backward_loss[m]["max_pixels_lost"]),
                str(int(backward_loss[m]["scans_with_loss"])),
            ]
            for m in method_order
            if m in backward_loss
        ],
    )

    _add_table(
        doc,
        "Table M4 - Paired forward/backward method agreement summary",
        [
            "Method",
            "Pairs",
            "Mean ratio (F/B)",
            "Median ratio (F/B)",
            f"Mean delta avg (F-B, {DISPLAY_MODULUS_UNIT})",
            "Mean delta n_valid (F-B)",
        ],
        [
            [
                method_labels.get(r["method"], r["method"]),
                str(_i(r.get("n_pairs"))),
                _fmt_ratio(_f(r.get("mean_ratio_avg"))),
                _fmt_ratio(_f(r.get("median_ratio_avg"))),
                _fmt_modulus_display(_f(r.get("mean_delta_avg"))),
                _fmt_num(_f(r.get("mean_delta_n_valid"))),
            ]
            for r in paired_rows
        ],
    )
    _add_table(
        doc,
        "Table M5 - Default processing parameters shared across modulus comparison methods",
            ["Parameter", "Default value", "Meaning"],
        [
            ["Metric", "Modulus", "Per-scan modulus summary from AFM modulus TIFFs."],
            [
                "Units",
                f"Source: {SOURCE_MODULUS_UNIT}; report display: {DISPLAY_MODULUS_UNIT}",
                "The pipeline summary files report modulus in source-tagged kPa; report tables/plots rescale those values for readability.",
            ],
            ["Scan size", "5 um x 5 um", "Per-scan image size."],
            ["Pixel grid", "512 x 512", "Nominal pixel resolution of each scan."],
            ["Baseline route", "modulus_gwy_stats / gwy_stats", "Gwyddion preprocessing + Gwyddion statistics."],
            ["Row/col metadata", "Parsed from filename", "Grid-position matching for paired and heatmap comparisons."],
        ],
    )
    _add_table(
        doc,
        "Table M6 - Method variation table",
        ["Method", "What changes relative to baseline", "n_valid effect", "Interpretation role"],
        [
            ["gwy_ops_py_stats", "Statistics engine changes from Gwyddion to Python after Gwyddion preprocessing.", "Usually no n_valid change.", "Separates preprocessing effects from stats-engine effects."],
            ["raw_minmax", "Uses raw export plus deterministic min/max validity window.", "Usually no or minimal n_valid change.", "Checks whether a simple raw-data validity rule reproduces baseline behavior."],
            ["raw_chauvenet", "Uses raw export plus Chauvenet outlier filtering.", "Moderate n_valid reduction where outliers exist.", "Tests sensitivity to probabilistic outlier rejection."],
            ["raw_three_sigma", "Uses raw export plus 3-sigma outlier filtering.", "Larger n_valid reduction where outliers exist.", "Tests sensitivity to sigma-based outlier rejection."],
            ["Forward/backward paired summary", "Compares matched scans by direction.", "Not a filtering route.", "Checks directional consistency for method selection."],
        ],
    )

    doc.add_paragraph(
        "Interpretive note: the paired forward/backward table is the most direct basis for the forward-only selection used "
        "in the Stage 1 topography workflow. Values with median ratio near 1.0 indicate close directional agreement when "
        "matching scans by grid position."
    )

    labels = [method_labels[m] for m in method_order]
    forward_delta_series = [
        [_f(r["delta_avg"]) * DISPLAY_MODULUS_SCALE for r in forward_long if r["method"] == m and _f(r["delta_avg"]) == _f(r["delta_avg"])]
        for m in method_order
    ]
    backward_delta_series = [
        [_f(r["delta_avg"]) * DISPLAY_MODULUS_SCALE for r in backward_long if r["method"] == m and _f(r["delta_avg"]) == _f(r["delta_avg"])]
        for m in method_order
    ]
    paired_ratio_series = [
        [_f(r["ratio_avg"]) for r in paired_long if r["method"] == m and _f(r["ratio_avg"]) == _f(r["ratio_avg"])]
        for m in method_order
        if any(pr["method"] == m for pr in paired_long)
    ]
    paired_ratio_labels = [method_labels[m] for m in method_order if any(pr["method"] == m for pr in paired_long)]

    forward_box = asset_dir / "forward_delta_avg_boxplot.png"
    backward_box = asset_dir / "backward_delta_avg_boxplot.png"
    paired_box = asset_dir / "paired_ratio_boxplot.png"
    forward_abs_bar = asset_dir / "forward_absolute_modulus_bar.png"
    backward_abs_bar = asset_dir / "backward_absolute_modulus_bar.png"
    forward_loss_bar = asset_dir / "forward_pixels_lost_bar.png"
    backward_loss_bar = asset_dir / "backward_pixels_lost_bar.png"
    forward_pct_loss_bar = asset_dir / "forward_percent_pixels_lost_bar.png"
    backward_pct_loss_bar = asset_dir / "backward_percent_pixels_lost_bar.png"
    combined_baseline_heatmap = asset_dir / "combined_forward_backward_baseline_heatmap.png"
    _plot_boxplot(
        forward_box,
        "Forward modulus comparison: delta avg vs baseline",
        labels,
        forward_delta_series,
        f"Delta avg ({DISPLAY_MODULUS_UNIT})",
    )
    _plot_boxplot(
        backward_box,
        "Backward modulus comparison: delta avg vs baseline",
        labels,
        backward_delta_series,
        f"Delta avg ({DISPLAY_MODULUS_UNIT})",
    )
    _plot_boxplot(paired_box, "Paired forward/backward ratio by method", paired_ratio_labels, paired_ratio_series, "Mean ratio (F/B)")
    _plot_bar_with_error(
        forward_abs_bar,
        "Forward absolute modulus by method",
        [method_labels[m] for m in method_order if m in forward_abs],
        [forward_abs[m]["mean_avg_value"] * DISPLAY_MODULUS_SCALE for m in method_order if m in forward_abs],
        [forward_abs[m]["se_avg_value"] * DISPLAY_MODULUS_SCALE for m in method_order if m in forward_abs],
        f"Mean modulus ({DISPLAY_MODULUS_UNIT})",
    )
    _plot_bar_with_error(
        backward_abs_bar,
        "Backward absolute modulus by method",
        [method_labels[m] for m in method_order if m in backward_abs],
        [backward_abs[m]["mean_avg_value"] * DISPLAY_MODULUS_SCALE for m in method_order if m in backward_abs],
        [backward_abs[m]["se_avg_value"] * DISPLAY_MODULUS_SCALE for m in method_order if m in backward_abs],
        f"Mean modulus ({DISPLAY_MODULUS_UNIT})",
    )
    _plot_bar_with_error(
        forward_loss_bar,
        "Forward pixels lost relative to baseline by method",
        [method_labels[m] for m in method_order if m in forward_loss],
        [forward_loss[m]["mean_pixels_lost"] for m in method_order if m in forward_loss],
        [forward_loss[m]["se_pixels_lost"] for m in method_order if m in forward_loss],
        "Pixels lost vs baseline",
    )
    _plot_bar_with_error(
        backward_loss_bar,
        "Backward pixels lost relative to baseline by method",
        [method_labels[m] for m in method_order if m in backward_loss],
        [backward_loss[m]["mean_pixels_lost"] for m in method_order if m in backward_loss],
        [backward_loss[m]["se_pixels_lost"] for m in method_order if m in backward_loss],
        "Pixels lost vs baseline",
    )
    _plot_bar_with_error(
        forward_pct_loss_bar,
        "Forward percent pixels lost relative to baseline by method",
        [method_labels[m] for m in method_order if m in forward_loss],
        [forward_loss[m]["mean_pct_lost"] for m in method_order if m in forward_loss],
        [forward_loss[m]["se_pct_lost"] for m in method_order if m in forward_loss],
        "Percent pixels lost vs baseline",
    )
    _plot_bar_with_error(
        backward_pct_loss_bar,
        "Backward percent pixels lost relative to baseline by method",
        [method_labels[m] for m in method_order if m in backward_loss],
        [backward_loss[m]["mean_pct_lost"] for m in method_order if m in backward_loss],
        [backward_loss[m]["se_pct_lost"] for m in method_order if m in backward_loss],
        "Percent pixels lost vs baseline",
    )
    _plot_combined_baseline_heatmap(combined_baseline_heatmap, paired_long)

    _add_figure(
        doc,
        "Figure M1 - Paired forward/backward modulus summary.",
        paired_dir / "plots" / "fwd_bwd_summary.png",
    )
    _add_figure(
        doc,
        "Figure M2 - Forward modulus method comparison (grouped bar plot).",
        forward_compare_dir / "plots" / "avg_grouped_bar.png",
    )
    _add_figure(
        doc,
        "Figure M3 - Backward modulus method comparison (grouped bar plot).",
        backward_compare_dir / "plots" / "avg_grouped_bar.png",
    )
    _add_figure(
        doc,
        f"Figure M3a - Forward modulus box plot of delta avg vs baseline by method ({DISPLAY_MODULUS_UNIT} display).",
        forward_box,
    )
    _add_figure(
        doc,
        f"Figure M3b - Backward modulus box plot of delta avg vs baseline by method ({DISPLAY_MODULUS_UNIT} display).",
        backward_box,
    )
    _add_figure(
        doc,
        "Figure M3c - Paired forward/backward ratio box plot by method.",
        paired_box,
    )
    _add_figure(
        doc,
        f"Figure M3d - Forward absolute modulus by method ({DISPLAY_MODULUS_UNIT} display). Bars show mean modulus across scans; error bars show standard error across scans.",
        forward_abs_bar,
    )
    _add_figure(
        doc,
        f"Figure M3e - Backward absolute modulus by method ({DISPLAY_MODULUS_UNIT} display). Bars show mean modulus across scans; error bars show standard error across scans.",
        backward_abs_bar,
    )
    _add_figure(
        doc,
        "Figure M3f - Forward pixels lost relative to baseline by method. Bars show mean pixels lost; error bars show standard error across scans.",
        forward_loss_bar,
    )
    _add_figure(
        doc,
        "Figure M3g - Backward pixels lost relative to baseline by method. Bars show mean pixels lost; error bars show standard error across scans.",
        backward_loss_bar,
    )
    _add_figure(
        doc,
        "Figure M3h - Forward percent pixels lost relative to baseline by method. Bars show mean percent loss; error bars show standard error across scans.",
        forward_pct_loss_bar,
    )
    _add_figure(
        doc,
        "Figure M3i - Backward percent pixels lost relative to baseline by method. Bars show mean percent loss; error bars show standard error across scans.",
        backward_pct_loss_bar,
    )
    _add_figure(
        doc,
        "Figure M4 - Forward baseline modulus heatmap.",
        forward_compare_dir / "plots" / "heatmap_two_panel__baseline.png",
    )
    _add_figure(
        doc,
        "Figure M5 - Backward baseline modulus heatmap.",
        backward_compare_dir / "plots" / "heatmap_two_panel__baseline.png",
    )
    _add_figure(
        doc,
        f"Figure M5a - Combined forward/backward baseline modulus heatmap ({DISPLAY_MODULUS_UNIT} display). Each cell is the mean of the paired forward and backward baseline values at the same grid position.",
        combined_baseline_heatmap,
    )

    _add_heading_paragraph(doc, "Forward delta-versus-baseline heatmaps")
    for method in ["gwy_ops_py_stats", "raw_minmax", "raw_chauvenet", "raw_three_sigma"]:
        _add_figure(
            doc,
            f"Forward delta heatmap - {method_labels.get(method, method)}.",
            forward_compare_dir / "plots" / f"heatmap_two_panel__delta_vs_baseline__{method}.png",
        )

    _add_heading_paragraph(doc, "Backward delta-versus-baseline heatmaps")
    for method in ["gwy_ops_py_stats", "raw_minmax", "raw_chauvenet", "raw_three_sigma"]:
        _add_figure(
            doc,
            f"Backward delta heatmap - {method_labels.get(method, method)}.",
            backward_compare_dir / "plots" / f"heatmap_two_panel__delta_vs_baseline__{method}.png",
        )

    doc.add_paragraph(
        "Use in Chapter 6: this report supports the baseline PEGDA validation section by showing that the modulus workflow "
        "was internally consistent across route variants and that forward/backward directional differences were small enough "
        "to justify restricting the present Stage 1 topography workflow to forward scans."
    )
    _add_heading_paragraph(doc, "Scope and next-step note")
    doc.add_paragraph(
        "This modulus baseline comparison was performed on a single PEGDA sample and one fracture-surface side. It is "
        "therefore a method-validation artifact rather than a full population-level modulus study. A more complete "
        "directional and route-consistency study should be repeated across additional grouped samples, separated by side, "
        "before broad generalization. The same comparison framework should later be extended to composite PEGDA-SiNP "
        "systems and Li-containing systems so that method consistency can be evaluated under the full set of material "
        "conditions used in the thesis."
    )

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)


def main() -> None:
    ap = argparse.ArgumentParser(description="Generate a compact modulus baseline validation report.")
    ap.add_argument(
        "--forward-compare-dir",
        default="out/method_compare/compare_20260223_173256",
    )
    ap.add_argument(
        "--backward-compare-dir",
        default="out/method_compare/compare_20260223_173335",
    )
    ap.add_argument(
        "--paired-dir",
        default="out/method_compare/fwd_bwd_20260223_173826",
    )
    ap.add_argument(
        "--forward-baseline-summary",
        default="out/method_compare/compare_inputs_20260223_173037/forward/gwy_stats/summary.csv",
    )
    ap.add_argument(
        "--backward-baseline-summary",
        default="out/method_compare/compare_inputs_20260223_173037/backward/gwy_stats/summary.csv",
    )
    ap.add_argument(
        "--output-docx",
        default=r"C:\Users\Conor O'Brien\Dropbox\03_AML\00 IN-BOX\AFM Topo Particle processing OUT\Modulus_Baseline_Validation_Report_20260306.docx",
    )
    args = ap.parse_args()
    build_report(
        forward_compare_dir=Path(args.forward_compare_dir),
        backward_compare_dir=Path(args.backward_compare_dir),
        paired_dir=Path(args.paired_dir),
        forward_baseline_summary=Path(args.forward_baseline_summary),
        backward_baseline_summary=Path(args.backward_baseline_summary),
        output_docx=Path(args.output_docx),
    )


if __name__ == "__main__":
    main()

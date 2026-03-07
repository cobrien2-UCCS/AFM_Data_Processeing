import csv
import math
import re
import statistics as stats
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches
from topo_report_synthesis import build_bundle, write_bundle

OUT_BASE = Path(r"C:\Users\Conor O'Brien\Dropbox\03_AML\00 IN-BOX\AFM Topo Particle processing OUT")
REPORT_PATH = OUT_BASE / "topo_particle_report_draft.docx"
DATA_GROUPED = Path("docs/File Locations for Data Grouped.txt")

BASELINE_JOB = "particle_forward_medianbg_mean"
SYSTEM_SINP = "PEGDA_SiNP"
TARGET_ISOLATED = 30
ISOLATION_DIST_NM = 900.0
DIAMETER_FILTER_NM = "350-550 nm (current runs)"
RISK_SCAN_COUNTS = [5, 10, 20, 30, 40]
JOB_ORDER = [
    "particle_forward_medianbg_mean",
    "particle_forward_medianbg_fixed0",
    "particle_forward_medianbg_p95",
    "particle_forward_medianbg_max_fixed0_p95",
    "particle_forward_flatten_mean",
    "particle_forward_flatten_fixed0",
    "particle_forward_flatten_p95",
    "particle_forward_flatten_max_fixed0_p95",
]


def read_csv_dicts(path):
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        return list(reader)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, val in enumerate(row):
            cells[idx].text = str(val)
    return table


def add_picture_if_exists(doc, path, width_in=5.5, caption=None):
    if path.exists():
        doc.add_picture(str(path), width=Inches(width_in))
        if caption:
            doc.add_paragraph(caption)
        else:
            doc.add_paragraph(f"Source: {path}")
        return True
    doc.add_paragraph(f"(Missing figure: {path.name})")
    return False


def add_job_histograms(doc, input_bases, job):
    bases = input_bases or [OUT_BASE]
    for base in bases:
        job_dir = base / "summary_outputs" / "job_hists" / job
        if not job_dir.exists():
            continue
        doc.add_paragraph(f"Job histograms for {job} ({base.name}):")
        for name, label in [
            ("hist_kept_counts.png", "Kept counts"),
            ("hist_raw_counts.png", "Raw counts"),
            ("hist_isolated_counts.png", "Isolated counts"),
        ]:
            path = job_dir / name
            if path.exists():
                doc.add_paragraph(label)
                add_picture_if_exists(doc, path, width_in=4.5)


def describe_job(job):
    if "medianbg" in job:
        preprocess = "align_rows x2 + plane_level + median_level + median(3)"
    elif "flatten" in job:
        preprocess = "align_rows x2 + plane_level + flatten_base + median(3)"
    else:
        preprocess = "custom/other"
    if job.endswith("_mean"):
        threshold = "mean"
    elif job.endswith("_fixed0"):
        threshold = "fixed 0.0"
    elif job.endswith("_p95"):
        threshold = "percentile 95"
    elif "max_fixed0_p95" in job:
        threshold = "max(mean, fixed0, p95)"
    else:
        threshold = "custom/other"
    return preprocess, threshold


def _load_fit_rows(input_bases):
    bases = input_bases or [OUT_BASE]
    summary_rows = []
    curve_rows = []
    for base in bases:
        fit_dir = base / "summary_outputs" / "fits"
        summary_rows.extend(read_csv_dicts(fit_dir / "fit_summary.csv"))
        curve_rows.extend(read_csv_dicts(fit_dir / "fit_risk_curves.csv"))
    return summary_rows, curve_rows


def _risk_prob_at(curve_rows, match, n_scans):
    for r in curve_rows:
        if all(r.get(k, "") == v for k, v in match.items()):
            if int(float(r.get("n_scans", 0))) == n_scans:
                return r.get("success_prob", "")
    return ""


def _pretty_aggregate_plot_name(path):
    stem = path.stem
    m = re.search(r"risk_aggregate_(.+)_poisson$", stem)
    label = m.group(1) if m else stem
    label = label.replace("wt_percent_", "").replace("_", " ")
    label = re.sub(r"\s+", " ", label).strip()
    label = re.sub(r"\bAll scrape states\b", "All scrape states", label, flags=re.I)
    label = re.sub(r"\bNon[- ]scraped\b", "Non-scraped", label, flags=re.I)
    label = re.sub(r"\bScraped\b", "Scraped", label, flags=re.I)
    label = re.sub(r"\b10\b", "10%", label)
    label = re.sub(r"\b25\b", "25%", label)
    label = re.sub(r"^(10%|25%)\\s+All scrape states$", r"\\1 | All scrape states", label)
    label = re.sub(r"^(10%|25%)\\s+Non-scraped$", r"\\1 | Non-scraped", label)
    label = re.sub(r"^(10%|25%)\\s+Scraped$", r"\\1 | Scraped", label)
    return label.strip()


def collect_debug_stats(bases):
    rows = []
    for base in bases:
        for debug_dir in base.rglob("debug"):
            if not debug_dir.is_dir():
                continue
            counts = {}
            total = 0
            for p in debug_dir.rglob("*"):
                if not p.is_file():
                    continue
                total += 1
                ext = p.suffix.lower() or "no_ext"
                counts[ext] = counts.get(ext, 0) + 1
            rows.append({
                "path": str(debug_dir),
                "total": total,
                "counts": counts,
            })
    return rows


def parse_data_grouped(path):
    groups = []
    current_system = None
    current_label = None
    for raw in path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# PEGDA Only"):
            current_system = "PEGDA"
            current_label = None
            continue
        if line.startswith("# PEGDA SiNP"):
            current_system = "PEGDA_SiNP"
            current_label = None
            continue
        if line.startswith("## "):
            current_label = line[3:].strip()
            groups.append({"system": current_system, "label": current_label, "roots": []})
            continue
        if line.startswith("#"):
            continue
        if line.startswith("C:\\") and groups:
            groups[-1]["roots"].append(line)
    return groups


def _norm_key(value):
    if not value:
        return ""
    return re.sub(r"[^a-z0-9]+", "", value.lower())


def build_sample_group_map(groups):
    mapping = {}
    norm_map = {}
    order = []
    root_norms = []
    for g in groups:
        label = g["label"]
        order.append(label)
        for root in g["roots"]:
            sample = Path(root).name
            mapping[sample] = g
            norm = _norm_key(sample)
            if norm:
                norm_map[norm] = g
            stripped = sample.rstrip("-_ ")
            if stripped and stripped not in mapping:
                mapping[stripped] = g
                norm_map[_norm_key(stripped)] = g
            root_norms.append((_norm_key(root), g))
    return mapping, norm_map, order, root_norms


def classify_sample(sample, mapping, norm_map, root_norms=None):
    if sample in mapping:
        return mapping[sample], "exact"
    norm = _norm_key(sample)
    if norm in norm_map:
        return norm_map[norm], "normalized"
    if root_norms:
        for root_norm, grp in root_norms:
            if root_norm and norm and norm in root_norm:
                return grp, "root"
    best = None
    best_len = 0
    for norm_key, grp in norm_map.items():
        if not norm_key:
            continue
        if norm_key in norm or norm in norm_key:
            if len(norm_key) > best_len:
                best_len = len(norm_key)
                best = grp
    return best, "fuzzy" if best else (None, None)


def parse_scan_id(source_file):
    m = re.search(r"LOC_RC(?P<row>\d{3})(?P<col>\d{3})", source_file)
    if not m:
        return Path(source_file).stem
    return "R{}-C{}".format(m.group("row"), m.group("col"))


def parse_specimen_metadata(name):
    if not name:
        return {}
    # Use specimen ID portion before _Sam if present
    specimen = name
    m = re.search(r"^(?P<spec>[^_]+)", name)
    if m:
        specimen = m.group("spec")
    # Coating metadata is not encoded for this project; assume none unless explicitly provided elsewhere.
    meta = {"polymer": "", "tpo": "", "sinp": "", "litfsi": "", "coating": "none"}
    m = re.search(r"^(?P<polymer>[A-Za-z]+)(?P<tpo>\\d+)?TPO(?P<sinp>\\d+)?SiNP", specimen, re.I)
    if m:
        meta["polymer"] = m.group("polymer") or ""
        tpo = m.group("tpo") or ""
        sinp = m.group("sinp") or ""
        meta["tpo"] = tpo.lstrip("0") or tpo
        meta["sinp"] = sinp.lstrip("0") or sinp
    m = re.search(r"LiTFSI(?P<litfsi>\\d+)", specimen, re.I)
    if m:
        meta["litfsi"] = m.group("litfsi").lstrip("0") or m.group("litfsi")
    return meta


def format_metadata(meta):
    parts = []
    if meta.get("polymer"):
        parts.append(f"Polymer={meta['polymer']}")
    if meta.get("tpo"):
        parts.append(f"TPO={meta['tpo']}%")
    if meta.get("sinp"):
        parts.append(f"SiNP={meta['sinp']}%")
    if meta.get("litfsi"):
        parts.append(f"LiTFSI={meta['litfsi']}%")
    if meta.get("coating"):
        parts.append(f"Coating={meta['coating']}")
    return ", ".join(parts) if parts else "No coded metadata in filename"


def wt_percent_from_label(label):
    if not label:
        return ""
    m = re.search(r"(10|25)%", label)
    return f"{m.group(1)}%" if m else ""


def wt_percent_from_sample(sample):
    if not sample:
        return ""
    m = re.search(r"tpo(10|25)sinp", sample, re.I)
    return f"{m.group(1)}%" if m else ""


def scraped_status_from_label(label):
    if not label:
        return ""
    if re.search(r"\bnon[- ]scraped\b", label, re.I):
        return "Non Scraped"
    if re.search(r"\bscraped\b", label, re.I):
        return "Scraped"
    return ""


def _find_plot_paths(input_bases, system, sample, job):
    bases = input_bases or [OUT_BASE]
    for base in bases:
        plot_dir = base / system / sample / "summary_outputs"
        count_plot = plot_dir / f"fig_particle_count_grid_{job}.png"
        iso_plot = plot_dir / f"fig_isolated_count_grid_{job}.png"
        if count_plot.exists() or iso_plot.exists():
            return count_plot, iso_plot
    return None, None


def parse_args():
    import argparse
    ap = argparse.ArgumentParser(description="Generate topo particle Word report.")
    ap.add_argument("--out-base", default="", help="Override output base directory.")
    ap.add_argument("--out-base-list", default="", help="Semicolon-separated input roots for a combined report.")
    ap.add_argument("--report-path", default="", help="Override report output path.")
    return ap.parse_args()


def main():
    args = parse_args()
    global OUT_BASE, REPORT_PATH
    input_bases = []
    if args.out_base_list:
        input_bases = [Path(p.strip()) for p in args.out_base_list.split(";") if p.strip()]
    if args.out_base:
        OUT_BASE = Path(args.out_base)
    elif input_bases:
        OUT_BASE = input_bases[0]
    if args.report_path:
        REPORT_PATH = Path(args.report_path)
    else:
        REPORT_PATH = OUT_BASE / "topo_particle_report_draft.docx"

    synthesis_bundle = None
    if len(input_bases) >= 2:
        wt10_root = None
        wt25_root = None
        for base in input_bases:
            label = base.name.lower()
            if "wt10" in label:
                wt10_root = base
            elif "wt25" in label:
                wt25_root = base
        if wt10_root and wt25_root:
            try:
                synthesis_bundle = build_bundle(wt10_root, wt25_root)
                write_bundle(synthesis_bundle, REPORT_PATH.parent)
            except Exception as exc:
                print(f"WARN: failed to write topo synthesis bundle: {exc}")

    doc = Document()
    job_order = JOB_ORDER

    doc.add_heading("Topo Particle Summary Report (Draft)", level=0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(
        "Purpose: quantify particle counts and sizes from AFM topography scans, compare processing methods, "
        "and estimate scan counts needed to achieve a statistically reliable number of isolated particles."
    )
    if input_bases:
        doc.add_paragraph("Output roots:")
        for base in input_bases:
            doc.add_paragraph(str(base), style="List Bullet")
    else:
        doc.add_paragraph(f"Output root: {OUT_BASE}")
    doc.add_paragraph(f"Baseline job: {BASELINE_JOB}")
    doc.add_paragraph("Particle coating: none (no coating metadata in filenames).")
    doc.add_paragraph("Map/grid area: 50 um x 50 um (21 x 21 grid, 5% overlap per scan)")
    if DATA_GROUPED.exists():
        doc.add_paragraph(f"Grouping source: {DATA_GROUPED}")

    doc.add_heading("Definitions and Processing Types", level=1)
    doc.add_paragraph(
        "Map = the full 50 um x 50 um grid footprint for a sample (21 x 21 scan positions). "
        "Scan = a single 5 um x 5 um AFM image at one grid index (Row/Col)."
    )
    doc.add_paragraph(
        "Processing steps are implemented using Gwyddion/pygwy operations; "
        "see the Gwyddion documentation for detailed descriptions of each operation."
    )
    doc.add_paragraph(
        "Baseline processing (per scan): alignment/row correction, background leveling (median), "
        "particle detection with mean-threshold mask, edge exclusion, diameter filter, and "
        "isolation filter based on center-to-center distance."
    )
    doc.add_paragraph(
        "Grid plots: missing scans are shown in gray; scans that exist but were filtered to zero "
        "particles are outlined in red; numbers overlay the raw particle count per scan."
    )
    doc.add_paragraph(
        "Grid scales: kept/isolated grids use 0-15 particles per scan; raw grids use 0-60 particles per scan."
    )
    doc.add_paragraph(
        "Unclassified = sample folder name did not match any grouped input roots in the "
        "File Locations for Data Grouped list."
    )
    doc.add_heading("Fit Models and Risk Equations", level=2)
    doc.add_paragraph(
        "Definitions: X_i = isolated particle count in scan i; S_n = sum of counts over n scans; "
        "T = target isolated particle count."
    )
    doc.add_paragraph(
        "Poisson: X_i ~ Poisson(lambda). P(X=k)=exp(-lambda)*lambda^k/k!. "
        "S_n ~ Poisson(n*lambda). P(S_n >= T) = 1 - sum_{k=0..T-1} P(S_n=k)."
    )
    doc.add_paragraph(
        "Negative Binomial (over-dispersed): Var(X)=mu+mu^2/r, r=mu^2/(Var-mu), p=r/(r+mu). "
        "X ~ NB(r,p). S_n ~ NB(n*r,p)."
    )
    doc.add_paragraph(
        "Zero-inflated NB: P(X=0)=pi+(1-pi)*NB(0;r,p); P(X=k)=(1-pi)*NB(k;r,p), k>0."
    )
    doc.add_paragraph(
        "Required scans: smallest n such that P(S_n >= T) >= q (q = 90%, 95%, 99%)."
    )
    doc.add_paragraph(
        "Uncertainty band: bootstrap the per-scan counts, recompute risk curves, and use "
        "5th/95th percentiles at each n."
    )
    doc.add_paragraph(
        "Method sensitivity: compare per-method histograms and fitted risk curves; compute "
        "histogram distances (JS divergence, L1, Wasserstein-1) and variance across methods "
        "for mean/variance/zero-rate."
    )
    doc.add_heading("Baseline Rationale", level=2)
    doc.add_paragraph(
        "Baseline is `particle_forward_medianbg_mean` (median-bg preprocessing + mean threshold). "
        "This is used as the reference method for percent-difference comparisons against other methods."
    )
    doc.add_paragraph("Processing methods (explicit):")
    method_rows = []
    for job in JOB_ORDER:
        preprocess, threshold = describe_job(job)
        method_rows.append([
            job,
            preprocess,
            threshold,
            ISOLATION_DIST_NM,
            DIAMETER_FILTER_NM,
        ])
    add_table(
        doc,
        ["Job", "Preprocess", "Threshold", "Isolation (nm)", "Diameter filter"],
        method_rows,
    )

    doc.add_heading("1. Scan Inventory", level=1)
    inv_rows = read_csv_dicts(OUT_BASE / "scan_inventory.csv")
    groups = parse_data_grouped(DATA_GROUPED) if DATA_GROUPED.exists() else []
    sample_to_group, norm_group_map, group_order, root_norms = (
        build_sample_group_map(groups) if groups else ({}, {}, [], [])
    )
    inv_by_root = {}
    inv_by_root_norm = {}
    inv_jsons = []
    if input_bases:
        inv_jsons = [b / "scan_inventory.json" for b in input_bases]
    else:
        inv_jsons = [OUT_BASE / "scan_inventory.json"]
    for inv_json in inv_jsons:
        if not inv_json.exists():
            continue
        try:
            import json
            for r in json.loads(inv_json.read_text(encoding="utf-8")):
                root = r.get("input_root")
                if root:
                    inv_by_root[root] = r
                    inv_by_root_norm[_norm_key(root)] = r
        except Exception:
            pass

    if inv_rows:
        headers = [
            "System",
            "SiNP wt%",
            "Scraped?",
            "Group",
            "Total scans",
            "Scan size (um x um)",
            "Pixel grid",
            "Resolution (nm/px)",
        ]
        rows = []
        if groups and inv_by_root:
            for g in groups:
                sys_name = g.get("system") or ""
                label = g.get("label") or ""
                wt = wt_percent_from_label(label)
                scraped = scraped_status_from_label(label)
                total = 0
                for root in g.get("roots") or []:
                    rec = inv_by_root.get(root) or inv_by_root_norm.get(_norm_key(root))
                    if rec:
                        total += int(rec.get("map_count", 0))
                total_display = total if total > 0 else "0 (not processed)"
                rows.append([
                    "PEGDA" if sys_name == "PEGDA" else "PEGDA-SiNP",
                    wt,
                    scraped,
                    label,
                    total_display,
                    "5 x 5",
                    "512 x 512",
                    "9.7656",
                ])
        else:
            for row in inv_rows:
                rows.append([
                    row.get("system", ""),
                    "",
                    "",
                    "All",
                    row.get("total_maps", ""),
                    f"{row.get('scan_um_x','')} x {row.get('scan_um_y','')}",
                    f"{row.get('grid_x','')} x {row.get('grid_y','')}",
                    row.get("resolution_nm_per_px", ""),
                ])
        add_table(doc, headers, rows)
    else:
        doc.add_paragraph("scan_inventory.csv not found.")

    doc.add_heading("2. Particle Count Data (PEGDA-SiNP)", level=1)
    fit_summary_rows, fit_curve_rows = _load_fit_rows(input_bases)
    if fit_summary_rows:
        doc.add_paragraph("Scan sufficiency and zero-rate summary (all jobs).")
        doc.add_paragraph("Scraped vs non-scraped are aggregated for distribution fits in this report.")
        doc.add_paragraph(
            "In this section, the modeled quantity is the isolated-particle count per scan. "
            "A probability such as P(total >= 30) @ 40 scans means the fitted model predicts the probability "
            "of reaching at least 30 isolated particles in total after 40 scans, not merely the probability of seeing one particle."
        )
        fit_headers = ["SiNP wt%", "Scraped?", "Job", "Model", "Zero-rate (obs)", "P0 (model)"]
        for n in RISK_SCAN_COUNTS:
            fit_headers.append(f"P(total >= {TARGET_ISOLATED}) @ {n} scans")
        fit_rows = []
        for r in fit_summary_rows:
            if r.get("count_field") != "count_isolated":
                continue
            model = r.get("count_model", "")
            wt = r.get("wt_percent", "")
            scraped = r.get("scraped", "")
            job = r.get("job", "")
            match = {
                "count_model": model,
                "job": job,
                "wt_percent": wt,
                "scraped": scraped,
                "count_field": "count_isolated",
            }
            row = [
                wt,
                scraped,
                job,
                model,
                r.get("zero_rate_obs", ""),
                r.get("p0_model", ""),
            ]
            for n in RISK_SCAN_COUNTS:
                row.append(_risk_prob_at(fit_curve_rows, match, n))
            fit_rows.append(row)
        if fit_rows:
            add_table(doc, fit_headers, fit_rows)
            doc.add_paragraph(
                "P0 (model) is the fitted probability of zero isolated particles in a scan. "
                "Zero-rate (obs) is the empirical fraction of zero-count scans."
            )
            doc.add_paragraph(
                "The probability of at least one non-zero scan across n scans is approximately "
                "1 - (zero_rate_obs)^n (independence assumption)."
            )
            doc.add_paragraph(
                "For rows derived from isolated counts, `max_per_scan` in the fit summary refers to the maximum isolated-particle count observed in a single scan. "
                "This is different from the retained candidate-particle maximum reported elsewhere in the report."
            )
        # Aggregate uncertainty plots (Poisson)
        for base in (input_bases or [OUT_BASE]):
            fit_dir = base / "summary_outputs" / "fits"
            newest_by_label = {}
            for path in fit_dir.glob("risk_aggregate_*_poisson.png"):
                label = _pretty_aggregate_plot_name(path)
                prev = newest_by_label.get(label)
                if prev is None or path.stat().st_mtime > prev.stat().st_mtime:
                    newest_by_label[label] = path
            for label, path in sorted(newest_by_label.items()):
                doc.add_paragraph(f"Aggregate Poisson uncertainty: {label}")
                add_picture_if_exists(doc, path, width_in=5.5)
                doc.add_paragraph(
                    "Reading guide: each curve shows the modeled probability of reaching the target isolated-particle total as the number of scans increases. "
                    "Horizontal dashed lines mark probability thresholds (for example 90%, 95%, and 99%). "
                    "Vertical requirement markers indicate the scan count needed to cross the 95% threshold for each method."
                )

    count_rows = []
    if input_bases:
        for base in input_bases:
            count_rows.extend(read_csv_dicts(base / "particle_counts_by_map.csv"))
    else:
        count_rows = read_csv_dicts(OUT_BASE / "particle_counts_by_map.csv")
    if count_rows:
        methods_ran = sorted({r.get("job") for r in count_rows if r.get("job")})
        if methods_ran:
            doc.add_paragraph("Methods actually run in this report:")
            for m in methods_ran:
                doc.add_paragraph(m, style="List Bullet")
    base_rows = [
        r for r in count_rows
        if r.get("job") == BASELINE_JOB and r.get("system") == SYSTEM_SINP
    ]
    if base_rows:
        counts = [int(float(r.get("count_total", 0) or 0)) for r in base_rows]
        isolated = [int(float(r.get("count_isolated", 0) or 0)) for r in base_rows]
        raw_counts = [int(float(r.get("count_total_raw", 0) or 0)) for r in base_rows if r.get("count_total_raw")]
        filtered_counts = [int(float(r.get("count_total_filtered", 0) or 0)) for r in base_rows if r.get("count_total_filtered")]
        total_particles = sum(counts)
        mean_counts = stats.mean(counts)
        std_counts = stats.pstdev(counts) if len(counts) > 1 else 0.0
        min_counts = min(counts)
        max_counts = max(counts)
        mean_iso = stats.mean(isolated)
        std_iso = stats.pstdev(isolated) if len(isolated) > 1 else 0.0
        pct_iso = 100.0 * sum(1 for v in isolated if v > 0) / float(len(isolated))

        doc.add_paragraph("Aggregate (baseline job)")
        add_table(
            doc,
            ["Metric", "Value"],
            [
                ["Total particles detected", total_particles],
                ["Mean particles per scan", f"{mean_counts:.3f}"],
                ["Std. dev. particles per scan", f"{std_counts:.3f}"],
                ["Min particles per scan", min_counts],
                ["Max particles per scan", max_counts],
                ["Mean raw particles per scan (pre-filter)", f"{stats.mean(raw_counts):.3f}" if raw_counts else "n/a"],
                ["Mean filtered particles per scan", f"{stats.mean(filtered_counts):.3f}" if filtered_counts else "n/a"],
                ["Mean isolated particles per scan", f"{mean_iso:.3f}"],
                ["Std. dev. isolated particles per scan", f"{std_iso:.3f}"],
                ["% scans with >= 1 isolated particle", f"{pct_iso:.3f}"],
            ],
        )
        doc.add_paragraph(
            "Here, `Max particles per scan` refers to the retained candidate-particle count after masking, diameter filtering, and edge exclusion, but before the isolation subset is taken. "
            "`Mean isolated particles per scan` and its standard deviation describe the isolated subset directly."
        )

        doc.add_paragraph("Per-scan particle counts (baseline job), grouped by SiNP wt% and scraped status.")
        sinp_group_order = [g["label"] for g in groups if g.get("system") == SYSTEM_SINP] if groups else []
        rows_by_group = {label: [] for label in sinp_group_order} if sinp_group_order else {}
        unclassified = set()
        for r in base_rows:
            sample = r.get("sample", "")
            group, _match_mode = classify_sample(sample, sample_to_group, norm_group_map, root_norms)
            label = group.get("label", "Unclassified") if group else "Unclassified"
            wt = wt_percent_from_label(label) or wt_percent_from_sample(sample) or "Unknown"
            scan_id = parse_scan_id(r.get("source_file", ""))
            if label == "Unclassified":
                unclassified.add(sample)
            rows_by_group.setdefault(label, []).append([
                wt,
                label,
                sample,
                scan_id,
                r.get("count_total", ""),
                r.get("count_isolated", ""),
            ])

        doc.add_paragraph(
            "Counts are reported as: total kept particles (after diameter + edge filters) and isolated particles "
            "(subset meeting the isolation distance). Raw (pre-filter) counts are plotted separately when available."
        )
        doc.add_paragraph(
            "In the count histograms, frequency means the number of scans falling at each retained-particle count. "
            "In the grid heatmaps, the displayed values are mean counts per scan position across the grouped sample sets, so the maps summarize spatial pattern but do not directly show the between-sample standard deviation at each position."
        )

        rows_by_wt = {"10%": [], "25%": [], "Unknown": []}
        for label, group_rows in rows_by_group.items():
            for r in group_rows:
                wt = r[0]
                group_label = r[1]
                scraped = scraped_status_from_label(group_label)
                rows_by_wt.setdefault(wt, []).append([
                    wt,
                    scraped,
                    group_label,
                    r[2],
                    r[3],
                    r[4],
                    r[5],
                ])

        for wt in ["10%", "25%", "Unknown"]:
            wt_rows = rows_by_wt.get(wt) or []
            if not wt_rows:
                continue
            doc.add_heading(f"SiNP {wt}", level=3)
            add_table(
                doc,
                ["SiNP wt%", "Scraped?", "Group", "Sample", "Scan ID", "Particles (Kept)", "Particles (Isolated)"],
                wt_rows,
            )
            samples = sorted({r[3] for r in wt_rows})
            for sample in samples:
                count_plot, iso_plot = _find_plot_paths(input_bases, SYSTEM_SINP, sample, BASELINE_JOB)
                meta = parse_specimen_metadata(sample)
                meta_str = format_metadata(meta)
                raw_plot = None
                for base in (input_bases or [OUT_BASE]):
                    candidate = (
                        base / SYSTEM_SINP / sample / "summary_outputs" / f"fig_particle_count_raw_grid_{BASELINE_JOB}.png"
                    )
                    if candidate.exists():
                        raw_plot = candidate
                        break
                if (count_plot and count_plot.exists()) or (iso_plot and iso_plot.exists()) or raw_plot:
                    doc.add_paragraph(
                        f"Grid density maps for {sample} ({BASELINE_JOB})\n"
                        f"{meta_str}\n"
                        "Counts are normalized per scan area (particles/um^2) with fixed color range "
                        "equivalent to 0-15 particles per scan. Raw grids (pre-filter) use 0-60."
                    )
                    if count_plot:
                        add_picture_if_exists(doc, count_plot)
                    if iso_plot:
                        add_picture_if_exists(doc, iso_plot)
                    if raw_plot:
                        add_picture_if_exists(doc, raw_plot)

        if unclassified:
            doc.add_paragraph("Unclassified samples (did not match grouped file list):")
            for sample in sorted(unclassified):
                doc.add_paragraph(sample, style="List Bullet")
    else:
        doc.add_paragraph("No baseline rows found in particle_counts_by_map.csv.")

    doc.add_heading("3. Particle Size Distribution", level=1)
    stats_rows = []
    if input_bases:
        for base in input_bases:
            stats_rows.extend(read_csv_dicts(base / "particle_summary_stats.csv"))
    else:
        stats_rows = read_csv_dicts(OUT_BASE / "particle_summary_stats.csv")
    mean_d = next((r["value"] for r in stats_rows if r.get("metric") == "mean_diameter_nm"), None)
    std_d = next((r["value"] for r in stats_rows if r.get("metric") == "std_diameter_nm"), None)
    diam_min_vals = sorted({r.get("diam_min_nm") for r in base_rows if r.get("diam_min_nm")})
    diam_max_vals = sorted({r.get("diam_max_nm") for r in base_rows if r.get("diam_max_nm")})
    diam_used = ""
    if diam_min_vals and diam_max_vals:
        diam_used = f"{diam_min_vals[0]}-{diam_max_vals[-1]} nm (from data)"
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Mean particle diameter (nm)", mean_d or "n/a"],
            ["Std. dev. particle diameter (nm)", std_d or "n/a"],
            ["Diameter filter (data)", diam_used or "n/a"],
            ["Diameter filter (target)", DIAMETER_FILTER_NM],
        ],
    )
    doc.add_paragraph(
        "Note: diameter stats are based on available per-particle exports; the baseline job may not export per-particle CSVs."
    )
    add_picture_if_exists(doc, OUT_BASE / "fig_particle_diameter_hist.png")

    # Per-method diameter stats by wt%
    diam_rows = []
    if input_bases:
        for base in input_bases:
            diam_rows.extend(read_csv_dicts(base / "particle_diameter_stats_by_job_wt.csv"))
    else:
        diam_rows = read_csv_dicts(OUT_BASE / "particle_diameter_stats_by_job_wt.csv")
    if diam_rows:
        doc.add_paragraph("Particle diameter stats by method and wt% (kept particles).")
        diam_table = []
        for r in diam_rows:
            diam_table.append([
                r.get("wt_percent", ""),
                r.get("job", ""),
                r.get("count", ""),
                r.get("mean_diameter_nm", ""),
                r.get("std_diameter_nm", ""),
            ])
        add_table(
            doc,
            ["SiNP wt%", "Job", "Count", "Mean Diam (nm)", "Std Diam (nm)"],
            diam_table,
        )
        # Embed histograms if present
        for r in diam_rows:
            wt = (r.get("wt_percent") or "").replace("%", "")
            job = r.get("job", "")
            for base in (input_bases or [OUT_BASE]):
                path = base / "summary_outputs" / "diameter_by_job_wt" / f"hist_diameter_{job}_wt{wt}.png"
                if path.exists():
                    doc.add_paragraph(f"Diameter histogram: {job} (wt{wt}%)")
                    add_picture_if_exists(doc, path, width_in=4.5)
                    break

    doc.add_heading("4. Isolation Criteria", level=1)
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Isolation definition (center-to-center)", f">= {ISOLATION_DIST_NM} nm"],
        ],
    )
    add_picture_if_exists(doc, OUT_BASE / "fig_isolated_count_hist.png")

    doc.add_heading("5. Method Comparison (By wt%)", level=1)
    if not count_rows:
        doc.add_paragraph("No particle_counts_by_map.csv rows available for method comparison.")
    else:
        # Use per-scan rows to compute per-wt, per-job summary. This avoids key-collisions when combining multiple roots.
        rows_sinp = [r for r in count_rows if r.get("system") == SYSTEM_SINP]
        by_wt_job_rows = {}
        for r in rows_sinp:
            wt = r.get("wt_percent") or "Unknown"
            job = r.get("job") or ""
            if not job:
                continue
            by_wt_job_rows.setdefault((wt, job), []).append(r)

        # Determine baseline job per wt% (fallback to first present method).
        wt_to_baseline = {}
        for (wt, job), rows_list in by_wt_job_rows.items():
            if wt not in wt_to_baseline:
                wt_to_baseline[wt] = None
            if job == BASELINE_JOB:
                wt_to_baseline[wt] = BASELINE_JOB
        for wt in list(wt_to_baseline.keys()):
            if wt_to_baseline[wt]:
                continue
            for j in JOB_ORDER:
                if (wt, j) in by_wt_job_rows:
                    wt_to_baseline[wt] = j
                    break

        for wt in sorted({k[0] for k in by_wt_job_rows.keys()}):
            baseline = wt_to_baseline.get(wt) or BASELINE_JOB
            doc.add_heading(f"Method Comparison: {wt} (baseline={baseline})", level=3)

            baseline_rows = by_wt_job_rows.get((wt, baseline), [])
            base_iso_mean = stats.mean([int(float(r.get('count_isolated', 0) or 0)) for r in baseline_rows]) if baseline_rows else 0.0

            rows = []
            for job in JOB_ORDER:
                rows_list = by_wt_job_rows.get((wt, job), [])
                if not rows_list:
                    continue
                totals = [int(float(r.get("count_total", 0) or 0)) for r in rows_list]
                raw = [int(float(r.get("count_total_raw", 0) or 0)) for r in rows_list if r.get("count_total_raw") not in ("", None)]
                filt = [int(float(r.get("count_total_filtered", 0) or 0)) for r in rows_list if r.get("count_total_filtered") not in ("", None)]
                iso = [int(float(r.get("count_isolated", 0) or 0)) for r in rows_list]
                mean_iso = stats.mean(iso) if iso else 0.0
                std_iso = stats.pstdev(iso) if len(iso) > 1 else 0.0
                pct_iso = 100.0 * sum(1 for v in iso if v > 0) / float(len(iso)) if iso else 0.0
                pct_diff_iso = ((mean_iso - base_iso_mean) / base_iso_mean * 100.0) if base_iso_mean > 0 else None

                wt_num = None
                m = re.search(r"(\\d+)", wt)
                if m:
                    wt_num = float(m.group(1))
                norm_yield = (mean_iso / wt_num) if wt_num and wt_num > 0 else None
                load_cost = (wt_num / mean_iso) if wt_num and mean_iso > 0 else None

                rows.append([
                    wt,
                    job,
                    f"{stats.mean(raw):.3f}" if raw else "n/a",
                    f"{stats.mean(totals):.3f}" if totals else "n/a",
                    f"{stats.mean(filt):.3f}" if filt else "n/a",
                    f"{mean_iso:.3f}",
                    f"{std_iso:.3f}",
                    f"{pct_iso:.1f}%",
                    f"{pct_diff_iso:.1f}%" if pct_diff_iso is not None else "n/a",
                    f"{norm_yield:.4f}" if norm_yield is not None else "n/a",
                    f"{load_cost:.3f}" if load_cost is not None else "n/a",
                ])

            add_table(
                doc,
                [
                    "wt%",
                    "Job",
                    "Raw mean/scan",
                    "Kept mean/scan",
                    "Filtered mean/scan",
                    "Mean isolated/scan",
                    "Std isolated/scan",
                    "% scans w/ iso",
                    "% diff iso vs baseline",
                    "Normalized isolation yield (iso/scan/%wt)",
                    "Loading cost (%wt/(iso/scan))",
                ],
                rows,
            )
            doc.add_paragraph(
                "Table note: Normalized isolation yield = (mean isolated particles per scan) divided by SiNP wt%. "
                "Loading cost is the inverse: SiNP wt% divided by mean isolated particles per scan."
            )
    # Per-wt bar plots (if available)
    for base in (input_bases or [OUT_BASE]):
        plot_dir = base / "summary_outputs" / "compare_by_wt"
        if not plot_dir.exists():
            continue
        for wt in ("10pct", "25pct"):
            family = plot_dir / f"fig_method_family_counts_{wt}.png"
            kept = plot_dir / f"fig_particle_count_mean_by_job_{wt}.png"
            iso = plot_dir / f"fig_isolated_count_mean_by_job_{wt}.png"
            if family.exists():
                doc.add_paragraph(f"Processing-family comparison plot ({wt.replace('pct','%')})")
                add_picture_if_exists(doc, family, width_in=6.0)
            elif kept.exists() or iso.exists():
                doc.add_paragraph(f"Per-job comparison plots ({wt.replace('pct','%')})")
                if kept.exists():
                    add_picture_if_exists(doc, kept, width_in=5.0)
                if iso.exists():
                    add_picture_if_exists(doc, iso, width_in=5.0)
    # Per-wt comparison tables
    if count_rows:
        by_job_wt = {}
        for r in count_rows:
            wt = r.get("wt_percent") or "Unknown"
            job = r.get("job", "")
            if not job:
                continue
            by_job_wt.setdefault(wt, {}).setdefault(job, []).append(int(float(r.get("count_isolated", 0) or 0)))
        for wt, job_map in sorted(by_job_wt.items()):
            doc.add_heading(f"Method Comparison by wt% ({wt})", level=3)
            rows = []
            base_vals = job_map.get(BASELINE_JOB) or []
            base_mean = stats.mean(base_vals) if base_vals else 0.0
            wt_val = float(wt.replace("%", "")) if wt and wt.replace("%", "").isdigit() else None
            for job in job_order:
                vals = job_map.get(job) or []
                if not vals:
                    continue
                mean_iso = stats.mean(vals)
                std_iso = stats.pstdev(vals) if len(vals) > 1 else 0.0
                pct_diff = ((mean_iso - base_mean) / base_mean * 100.0) if base_mean > 0 else None
                ratio = (mean_iso / wt_val) if wt_val and wt_val > 0 else None
                inv_ratio = (wt_val / mean_iso) if mean_iso and mean_iso > 0 and wt_val else None
                rows.append([
                    job,
                    f"{mean_iso:.3f}",
                    f"{std_iso:.3f}",
                    f"{pct_diff:.1f}%" if pct_diff is not None else "n/a",
                    f"{ratio:.4f}" if ratio is not None else "n/a",
                    f"{inv_ratio:.4f}" if inv_ratio is not None else "n/a",
                ])
            add_table(
                doc,
                [
                    "Job",
                    "Mean Isolated/Scan",
                    "Std Isolated/Scan",
                    "% diff vs baseline",
                    "Mean Isolated/Scan per wt%",
                    "wt% per Isolated/Scan",
                ],
                rows,
            )
            doc.add_paragraph(
                "Footnote: Normalized isolation yield = (mean isolated/scan) / (SiNP wt%). "
                "Loading cost = (SiNP wt%) / (mean isolated/scan). "
                "Units: isolated·scan^-1·wt%^-1 and wt%·scan·isolated^-1."
            )
    doc.add_paragraph(
        "Per-job histograms (kept/raw/isolated) are embedded below and also exported under "
        "summary_outputs/job_hists/<job>/ in each output root."
    )
    doc.add_paragraph(
        "Box plots: center line = median; box = interquartile range (Q1-Q3); whiskers extend to "
        "1.5×IQR; points beyond whiskers are outliers."
    )
    # Method histogram sensitivity tables (from fit outputs)
    if fit_summary_rows:
        doc.add_paragraph("Method histogram sensitivity (from fit outputs).")
        method_dist = []
        method_var = []
        for base in (input_bases or [OUT_BASE]):
            fit_dir = base / "summary_outputs" / "fits"
            method_dist.extend(read_csv_dicts(fit_dir / "method_histogram_distances.csv"))
            method_var.extend(read_csv_dicts(fit_dir / "method_variance_summary.csv"))
        if method_dist:
            method_dist_sorted = sorted(method_dist, key=lambda r: float(r.get("js_divergence", 0) or 0), reverse=True)
            rows = []
            for r in method_dist_sorted[:10]:
                rows.append([
                    r.get("job_a", ""),
                    r.get("job_b", ""),
                    r.get("js_divergence", ""),
                    r.get("l1_distance", ""),
                    r.get("wasserstein1", ""),
                ])
            add_table(doc, ["Job A", "Job B", "JS Divergence", "L1 Distance", "Wasserstein-1"], rows)
        if method_var:
            rows = []
            for r in method_var:
                rows.append([
                    r.get("metric", ""),
                    r.get("job_count", ""),
                    r.get("mean", ""),
                    r.get("variance", ""),
                    r.get("std", ""),
                    r.get("min", ""),
                    r.get("max", ""),
                ])
            add_table(
                doc,
                ["Metric", "Jobs", "Mean", "Variance", "Std", "Min", "Max"],
                rows,
            )
    present_jobs = set()
    if count_rows:
        present_jobs = {r.get("job") for r in count_rows if r.get("job")}
    for job in job_order:
        if present_jobs and job not in present_jobs:
            continue
        add_job_histograms(doc, input_bases, job)
    for base in (input_bases or [OUT_BASE]):
        p1 = base / "fig_particle_count_mean_by_job.png"
        p2 = base / "fig_isolated_count_mean_by_job.png"
        if p1.exists() or p2.exists():
            doc.add_paragraph(f"Per-job means ({base.name}):")
        add_picture_if_exists(doc, p1)
        add_picture_if_exists(doc, p2)

    doc.add_heading("5A. Heatmap Uncertainty Companions", level=1)
    doc.add_paragraph(
        "The mean heatmaps in this report are paired with standard-deviation companions. "
        "These maps show the between-sample spread at each scan position for the grouped datasets "
        "and provide the most direct uncertainty layer for the primary kept-count and isolated-count maps."
    )
    for base in (input_bases or [OUT_BASE]):
        combined_dir = base / "summary_outputs" / "combined"
        if not combined_dir.exists():
            continue
        for wt in ("10", "25"):
            doc.add_paragraph(f"Heatmap uncertainty companions ({wt} wt%)")
            for job in [BASELINE_JOB]:
                for stem, label in [
                    ("fig_particle_count_grid_std_wt{wt}_{job}.png", "Kept-count std map"),
                    ("fig_isolated_count_grid_std_wt{wt}_{job}.png", "Isolated-count std map"),
                ]:
                    path = combined_dir / stem.format(wt=wt, job=job)
                    if path.exists():
                        doc.add_paragraph(f"{label}: {job} ({wt} wt%)")
                        add_picture_if_exists(doc, path, width_in=5.0)

    doc.add_heading("6. Grain Summary (Selected Fields)", level=1)
    doc.add_paragraph("Grain exports are enabled in the particle modes for this report.")
    for base in (input_bases or [OUT_BASE]):
        trend_dir = base / "summary_outputs" / "grain_compare"
        if not trend_dir.exists():
            continue
        for name, label in [
            ("fig_grain_diameter_nm_kept_mean_by_job.png", "Grain diameter trend (kept): mean +/- std by job"),
            ("fig_grain_diameter_nm_kept_box_by_job.png", "Grain diameter distribution (kept): box plot by job"),
            ("fig_grain_diameter_nm_isolated_mean_by_job.png", "Grain diameter trend (isolated): mean +/- std by job"),
            ("fig_grain_diameter_nm_isolated_box_by_job.png", "Grain diameter distribution (isolated): box plot by job"),
        ]:
            path = trend_dir / name
            if path.exists():
                doc.add_paragraph(label)
                add_picture_if_exists(doc, path, width_in=5.5)
    grain_rows = []
    if input_bases:
        for base in input_bases:
            grain_rows.extend(read_csv_dicts(base / "grain_summary_by_job.csv"))
    else:
        grain_rows = read_csv_dicts(OUT_BASE / "grain_summary_by_job.csv")
    grain_map = {r.get("job"): r for r in grain_rows}
    grain_table_rows = []
    for job in job_order:
        r = grain_map.get(job)
        if not r:
            continue
        grain_table_rows.append([
            job,
            r.get("grain_total", ""),
            r.get("grain_kept", ""),
            r.get("grain_isolated", ""),
            r.get("kept_mean_diameter_nm", ""),
            r.get("kept_std_diameter_nm", ""),
            r.get("isolated_mean_diameter_nm", ""),
            r.get("isolated_std_diameter_nm", ""),
        ])
    add_table(
        doc,
        [
            "Job",
            "Grain Total",
            "Grain Kept",
            "Grain Isolated",
            "Kept Mean Diam (nm)",
            "Kept Std Diam (nm)",
            "Isolated Mean Diam (nm)",
            "Isolated Std Diam (nm)",
        ],
        grain_table_rows,
    )
    # Per-wt grain summary if available
    grain_sample_rows = []
    if input_bases:
        for base in input_bases:
            grain_sample_rows.extend(read_csv_dicts(base / "grain_summary_by_sample_job.csv"))
    else:
        grain_sample_rows = read_csv_dicts(OUT_BASE / "grain_summary_by_sample_job.csv")
    if grain_sample_rows:
        doc.add_paragraph("Grain summary by method and wt% (from sample-level grain stats).")
        by_wt_job = {}
        for r in grain_sample_rows:
            sample = r.get("sample", "")
            wt = wt_percent_from_sample(sample) or "Unknown"
            job = r.get("job", "")
            if not job:
                continue
            by_wt_job.setdefault((wt, job), []).append(r)
        rows = []
        for (wt, job), rows_list in sorted(by_wt_job.items()):
            totals = [float(r.get("grain_total", 0) or 0) for r in rows_list]
            kept = [float(r.get("grain_kept", 0) or 0) for r in rows_list]
            iso = [float(r.get("grain_isolated", 0) or 0) for r in rows_list]
            kept_mean = [float(r.get("kept_mean_diameter_nm", 0) or 0) for r in rows_list]
            kept_std = [float(r.get("kept_std_diameter_nm", 0) or 0) for r in rows_list]
            rows.append([
                wt,
                job,
                f"{stats.mean(totals):.1f}" if totals else "n/a",
                f"{stats.mean(kept):.1f}" if kept else "n/a",
                f"{stats.mean(iso):.1f}" if iso else "n/a",
                f"{stats.mean(kept_mean):.2f}" if kept_mean else "n/a",
                f"{stats.mean(kept_std):.2f}" if kept_std else "n/a",
            ])
        add_table(
            doc,
            ["SiNP wt%", "Job", "Grain Total (mean)", "Grain Kept (mean)", "Grain Isolated (mean)", "Kept Mean Diam (nm)", "Kept Std Diam (nm)"],
            rows,
        )
        # Optional: embed grain diameter histograms by job if available
        for base in (input_bases or [OUT_BASE]):
            plots_dir = base / "grain_plots"
            if not plots_dir.exists():
                continue
            for job in job_order:
                job_dir = plots_dir / job
                if not job_dir.exists():
                    continue
                kept_hist = job_dir / "grain_diameter_nm_kept.png"
                if kept_hist.exists():
                    doc.add_paragraph(f"Grain diameter histogram (kept): {job} ({base.name})")
                    add_picture_if_exists(doc, kept_hist, width_in=5.5)
    else:
        doc.add_paragraph(
            "Grain exports were not available for this run. "
            "Expected *_grains.csv outputs were missing; rerun pygwy with grain_export enabled "
            "and use_review_sample=false to populate grain tables."
        )

    doc.add_heading("7. Statistical Feasibility Statement", level=1)
    if count_rows:
        rows_sinp = [r for r in count_rows if r.get("system") == SYSTEM_SINP]
        by_wt_job = {}
        for r in rows_sinp:
            wt = r.get("wt_percent") or "Unknown"
            job = r.get("job", "")
            iso = int(float(r.get("count_isolated", 0) or 0))
            by_wt_job.setdefault(wt, {}).setdefault(job, []).append(iso)

        for wt, job_map in sorted(by_wt_job.items()):
            doc.add_heading(f"Feasibility by wt% ({wt})", level=3)
            base_vals = job_map.get(BASELINE_JOB) or []
            base_mean = stats.mean(base_vals) if base_vals else 0.0
            method_rows = []
            scans_needed_list = []
            mean_iso_list = []
            for job in JOB_ORDER:
                vals = job_map.get(job) or []
                if not vals:
                    continue
                mean_iso = stats.mean(vals)
                std_iso = stats.pstdev(vals) if len(vals) > 1 else 0.0
                scans_needed = math.ceil(TARGET_ISOLATED / mean_iso) if mean_iso > 0 else None
                pct_diff = ((mean_iso - base_mean) / base_mean * 100.0) if base_mean > 0 else None
                if scans_needed:
                    scans_needed_list.append(scans_needed)
                mean_iso_list.append(mean_iso)
                method_rows.append([
                    job,
                    f"{mean_iso:.3f}",
                    f"{std_iso:.3f}",
                    str(scans_needed) if scans_needed else "n/a",
                    f"{pct_diff:.1f}%" if pct_diff is not None else "n/a",
                ])
            add_table(
                doc,
                ["Job", "Mean isolated/scan", "Std isolated/scan", "Scans for ~30 isolated", "% diff vs baseline"],
                method_rows,
            )
            if scans_needed_list:
                mean_scans = stats.mean(scans_needed_list)
                std_scans = stats.pstdev(scans_needed_list) if len(scans_needed_list) > 1 else 0.0
            else:
                mean_scans = 0.0
                std_scans = 0.0
            if mean_iso_list:
                mean_iso_all = stats.mean(mean_iso_list)
                std_iso_all = stats.pstdev(mean_iso_list) if len(mean_iso_list) > 1 else 0.0
            else:
                mean_iso_all = 0.0
                std_iso_all = 0.0
            doc.add_paragraph(
                f"Across methods (this wt%): mean scans needed = {mean_scans:.2f} (std {std_scans:.2f}); "
                f"mean isolated/scan = {mean_iso_all:.3f} (std {std_iso_all:.3f})."
            )
        doc.add_paragraph(
            "Percent difference is computed as (method - baseline) / baseline * 100 using the baseline job mean isolated/scan (within each wt%)."
        )
    else:
        doc.add_paragraph("Feasibility statement could not be computed (no isolated count data).")

    doc.add_heading("8. Equations (Summary)", level=1)
    doc.add_paragraph(
        "Variables: N = number of scans, x_i = particle count for scan i, x_bar = mean of x_i, "
        "sigma = population std of x_i, T = target isolated particle count (30), "
        "N_needed = scans required to reach T in expectation."
    )
    doc.add_paragraph("Mean per scan: x_bar = (1/N) * sum(x_i)")
    doc.add_paragraph("Population std: sigma = sqrt((1/N) * sum((x_i - x_bar)^2))")
    doc.add_paragraph("Percent scans with isolated: 100 * sum(1[x_i > 0]) / N")
    doc.add_paragraph("Scans needed: N_needed = ceil(T / mean_isolated_per_scan), with T = 30")
    doc.add_paragraph("Expected isolated for N scans: E_total = N * mean_isolated_per_scan")
    doc.add_paragraph("Std of isolated total for N scans: sigma_total = sqrt(N) * std_isolated_per_scan")
    doc.add_paragraph("Resolution: 5000 nm / 512 = 9.7656 nm/px")

    doc.add_heading("9. LiTFSI Comparison", level=1)
    doc.add_paragraph("Not applicable (no PEGDA-LiTFSI-SiNP samples in this dataset).")

    doc.add_heading("10. Run Time and Debug Outputs", level=1)
    timing_rows = []
    bases_for_timing = input_bases or [OUT_BASE]
    for base in bases_for_timing:
        timing_path = base / "run_timing.json"
        if not timing_path.exists():
            continue
        try:
            import json
            info = json.loads(timing_path.read_text(encoding="utf-8"))
            total_s = float(info.get("total_seconds", 0.0))
            timing_rows.append([
                base.name,
                f"{total_s:.1f}",
                info.get("total_scans", ""),
                info.get("total_jobs", ""),
                info.get("roots_processed", ""),
                info.get("started", ""),
                info.get("finished", ""),
            ])
        except Exception:
            continue
    if timing_rows:
        add_table(
            doc,
            ["Output Root", "Total Seconds", "Total Scans", "Jobs", "Roots", "Start", "Finish"],
            timing_rows,
        )
    else:
        doc.add_paragraph("run_timing.json not found (time data not available for this run).")

    # Optional estimate note
    for base in bases_for_timing:
        estimate_path = base / "run_estimate.txt"
        if estimate_path.exists():
            doc.add_paragraph(f"Estimated run time note ({base.name}):")
            try:
                note_text = estimate_path.read_text(encoding="utf-8")
            except UnicodeDecodeError:
                note_text = estimate_path.read_text(encoding="cp1252", errors="replace")
            doc.add_paragraph(note_text)

    doc.add_paragraph(
        "Run-time drivers: number of scans x number of methods, Gwyddion preprocessing steps, "
        "particle detection + grain exports, and debug artifact saves (aligned/leveled/filtered/mask)."
    )

    debug_stats = collect_debug_stats(bases_for_timing)
    if debug_stats:
        doc.add_paragraph("Debug artifact summary (per debug directory):")
        debug_rows = []
        for row in debug_stats:
            counts = ", ".join([f"{k}:{v}" for k, v in sorted(row["counts"].items())])
            debug_rows.append([row["path"], row["total"], counts])
        add_table(doc, ["Debug Path", "Total Files", "By Extension"], debug_rows[:20])
        if len(debug_rows) > 20:
            doc.add_paragraph(f"... {len(debug_rows) - 20} more debug folders not shown.")
    else:
        doc.add_paragraph("No debug folders found under output roots.")

    # Per-method timing (if available)
    method_time_rows = []
    for base in bases_for_timing:
        timing_path = base / "run_timing.json"
        if not timing_path.exists():
            continue
        try:
            import json
            info = json.loads(timing_path.read_text(encoding="utf-8"))
            entries = info.get("entries") or []
            by_job = {}
            for e in entries:
                job = e.get("job")
                root = e.get("input_root", "")
                wt = "10%" if re.search(r"tpo10sinp", root, re.I) else ("25%" if re.search(r"tpo25sinp", root, re.I) else "Unknown")
                if not job:
                    continue
                key = (job, wt)
                by_job.setdefault(key, []).append(float(e.get("seconds", 0.0)))
            for (job, wt), vals in sorted(by_job.items()):
                if not vals:
                    continue
                method_time_rows.append([base.name, wt, job, f"{sum(vals):.1f}", f"{stats.mean(vals):.1f}", len(vals)])
        except Exception:
            continue
    if method_time_rows:
        doc.add_paragraph("Per-method timing summary (sum and mean over sample roots).")
        add_table(
            doc,
            ["Output Root", "SiNP wt%", "Job", "Total Seconds", "Mean Seconds", "Samples"],
            method_time_rows,
        )

    try:
        doc.save(str(REPORT_PATH))
        print(f"Wrote report to {REPORT_PATH}")
    except PermissionError:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        alt_path = OUT_BASE / f"topo_particle_report_draft_{ts}.docx"
        doc.save(str(alt_path))
        print(f"Report file was locked; wrote report to {alt_path}")


if __name__ == "__main__":
    main()

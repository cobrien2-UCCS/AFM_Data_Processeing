from __future__ import annotations

import argparse
import csv
import json
import math
from collections import defaultdict
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from statistics import mean, stdev

from docx import Document
from docx.shared import Inches
from topo_report_synthesis import build_bundle

PRIMARY_JOB = "particle_forward_medianbg_mean"
COMPARISON_JOB = "particle_forward_flatten_mean"
TARGET_ISOLATED = 30
CONFIDENCE = 0.95
SCAN_SIZE_UM = "5 x 5"
PIXEL_GRID = "512 x 512"
RESOLUTION_NM_PER_PIXEL = "9.77"
NOMINAL_GRID = "21 x 21"
DIAMETER_FILTER_NM = "350-550"
ISOLATION_DISTANCE_NM = "900"
ASSET_DIR = Path("docs/Thesis/generated_assets")
MODULUS_ASSET_DIR = Path(r"C:\Users\Conor O'Brien\Dropbox\03_AML\00 IN-BOX\AFM Topo Particle processing OUT\modulus_baseline_assets")
MODULUS_PAIRED_SUMMARY = Path("out/method_compare/fwd_bwd_20260223_173826/fwd_bwd_summary.csv")
MODULUS_FORWARD_COMPARE_DIR = Path("out/method_compare/compare_20260306_180402")
MODULUS_BACKWARD_COMPARE_DIR = Path("out/method_compare/compare_20260306_180452")
REP_IMAGE_DIR = Path(r"C:\Users\Conor O'Brien\Dropbox\03_AML\00 IN-BOX\AFM Topo Particle processing OUT\Particle Images for Thesis Figures\processed")

JOB_ORDER = [
    "particle_forward_medianbg_mean",
    "particle_forward_medianbg_fixed0",
    "particle_forward_medianbg_p95",
    "particle_forward_medianbg_max_fixed0_p95",
    "particle_forward_flatten_mean",
    "particle_forward_flatten_fixed0",
    "particle_forward_flatten_p95",
    "particle_forward_flatten_max_fixed0_p95",
]

JOB_LABELS = {
    "particle_forward_medianbg_mean": "Median BG + mean threshold",
    "particle_forward_medianbg_fixed0": "Median BG + fixed 0",
    "particle_forward_medianbg_p95": "Median BG + p95 threshold",
    "particle_forward_medianbg_max_fixed0_p95": "Median BG + fixed 0 + p95",
    "particle_forward_flatten_mean": "Flatten + mean threshold",
    "particle_forward_flatten_fixed0": "Flatten + fixed 0",
    "particle_forward_flatten_p95": "Flatten + p95 threshold",
    "particle_forward_flatten_max_fixed0_p95": "Flatten + fixed 0 + p95",
}
JOB_STYLES = {
    "particle_forward_medianbg_mean": {"color": "#1f77b4", "linestyle": "-", "marker": "o", "linewidth": 2.5},
    "particle_forward_medianbg_fixed0": {"color": "#2ca02c", "linestyle": "--", "marker": "s", "linewidth": 1.8},
    "particle_forward_medianbg_p95": {"color": "#d62728", "linestyle": "-.", "marker": "^", "linewidth": 1.8},
    "particle_forward_medianbg_max_fixed0_p95": {"color": "#9467bd", "linestyle": ":", "marker": "D", "linewidth": 1.8},
    "particle_forward_flatten_mean": {"color": "#ff7f0e", "linestyle": "-", "marker": "P", "linewidth": 2.5},
    "particle_forward_flatten_fixed0": {"color": "#8c564b", "linestyle": "--", "marker": "X", "linewidth": 1.8},
    "particle_forward_flatten_p95": {"color": "#e377c2", "linestyle": "-.", "marker": "v", "linewidth": 1.8},
    "particle_forward_flatten_max_fixed0_p95": {"color": "#7f7f7f", "linestyle": ":", "marker": "<", "linewidth": 1.8},
}


@dataclass
class JobStats:
    maps: int
    mean_count: float
    std_count: float
    min_count: int
    max_count: int
    zero_count_rate: float
    mean_isolated: float
    std_isolated: float
    min_isolated: int
    max_isolated: int
    zero_isolated_rate: float
    pct_with_isolated: float
    sample_means: dict[str, float]
    sample_stds: dict[str, float]
    sample_pct_nonzero: dict[str, float]
    sample_scan_counts: dict[str, int]


@dataclass
class FitStats:
    n_scans: int
    mean_per_scan: float
    variance_per_scan: float
    zero_rate_obs: float
    n_required_095: int


@dataclass
class GrainStats:
    grain_total: int
    grain_kept: int
    grain_isolated: int
    kept_mean_diameter_nm: float
    kept_std_diameter_nm: float
    isolated_mean_diameter_nm: float
    isolated_std_diameter_nm: float


@dataclass
class RootStats:
    label: str
    wt_percent: int
    root: Path
    samples: list[str]
    diameter_mean_nm: float
    diameter_std_nm: float
    method_stats: dict[str, JobStats]
    method_fits: dict[str, FitStats]
    grain_stats: dict[str, GrainStats]


def _read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def _f(value: str | None) -> float:
    return 0.0 if value in (None, "") else float(value)


def _i(value: str | None) -> int:
    return 0 if value in (None, "") else int(float(value))


def _pct(rate: float) -> str:
    return f"{100.0 * rate:.1f}%"


def _pm(value: float, std: float, unit: str = "") -> str:
    suffix = f" {unit}" if unit else ""
    return f"{value:.2f} +/- {std:.2f}{suffix}"


def _job_label(job: str) -> str:
    return JOB_LABELS.get(job, job)


def _job_stats(rows: list[dict[str, str]], job: str) -> JobStats:
    rows = [r for r in rows if r.get("job") == job]
    counts = [_i(r.get("count_total")) for r in rows]
    isolated = [_i(r.get("count_isolated")) for r in rows]
    sample_counts: dict[str, list[int]] = defaultdict(list)
    for row in rows:
        sample_counts[row.get("sample", "")].append(_i(row.get("count_isolated")))
    return JobStats(
        maps=len(rows),
        mean_count=mean(counts),
        std_count=stdev(counts),
        min_count=min(counts),
        max_count=max(counts),
        zero_count_rate=sum(1 for x in counts if x == 0) / len(counts),
        mean_isolated=mean(isolated),
        std_isolated=stdev(isolated),
        min_isolated=min(isolated),
        max_isolated=max(isolated),
        zero_isolated_rate=sum(1 for x in isolated if x == 0) / len(isolated),
        pct_with_isolated=100.0 * sum(1 for x in isolated if x > 0) / len(isolated),
        sample_means={k: mean(v) for k, v in sample_counts.items()},
        sample_stds={k: (stdev(v) if len(v) > 1 else 0.0) for k, v in sample_counts.items()},
        sample_pct_nonzero={k: 100.0 * sum(1 for x in v if x > 0) / len(v) for k, v in sample_counts.items()},
        sample_scan_counts={k: len(v) for k, v in sample_counts.items()},
    )


def _fit_map(path: Path) -> dict[str, FitStats]:
    out: dict[str, FitStats] = {}
    for row in _read_csv(path):
        if row.get("count_field") != "count_isolated" or row.get("count_model") != "poisson":
            continue
        out[row["job"]] = FitStats(
            n_scans=_i(row.get("n_scans")),
            mean_per_scan=_f(row.get("mean_per_scan")),
            variance_per_scan=_f(row.get("variance_per_scan")),
            zero_rate_obs=_f(row.get("zero_rate_obs")),
            n_required_095=_i(row.get("n_required_095")),
        )
    return out


def _grain_map(path: Path) -> dict[str, GrainStats]:
    out: dict[str, GrainStats] = {}
    for row in _read_csv(path):
        out[row["job"]] = GrainStats(
            grain_total=_i(row.get("grain_total")),
            grain_kept=_i(row.get("grain_kept")),
            grain_isolated=_i(row.get("grain_isolated")),
            kept_mean_diameter_nm=_f(row.get("kept_mean_diameter_nm")),
            kept_std_diameter_nm=_f(row.get("kept_std_diameter_nm")),
            isolated_mean_diameter_nm=_f(row.get("isolated_mean_diameter_nm")),
            isolated_std_diameter_nm=_f(row.get("isolated_std_diameter_nm")),
        )
    return out


def _summary_stats(path: Path) -> tuple[float, float]:
    values = {r["metric"]: r["value"] for r in _read_csv(path)}
    return float(values["mean_diameter_nm"]), float(values["std_diameter_nm"])


def _load_root(root: Path, wt_percent: int) -> RootStats:
    counts_rows = _read_csv(root / "particle_counts_by_map.csv")
    diam_mean, diam_std = _summary_stats(root / "particle_summary_stats.csv")
    method_stats = {job: _job_stats(counts_rows, job) for job in JOB_ORDER}
    return RootStats(
        label=f"{wt_percent} wt% SiNP",
        wt_percent=wt_percent,
        root=root,
        samples=sorted(method_stats[PRIMARY_JOB].sample_means),
        diameter_mean_nm=diam_mean,
        diameter_std_nm=diam_std,
        method_stats=method_stats,
        method_fits=_fit_map(root / "summary_outputs" / "fits" / "fit_summary.csv"),
        grain_stats=_grain_map(root / "grain_summary_by_job.csv"),
    )


def _figure(doc: Document, title: str, paths: list[Path], width: float = 5.9) -> None:
    existing = [p for p in paths if p.exists()]
    if not existing:
        doc.add_paragraph(f"[MISSING FIGURE] {title}")
        return
    for path in existing:
        doc.add_picture(str(path), width=Inches(width))
    doc.add_paragraph(title)
    for path in existing:
        doc.add_paragraph(f"Source: {path}")


def _combined_grid_paths(root: RootStats, prefix: str, job: str) -> list[Path]:
    base = root.root / "summary_outputs" / "combined"
    return [
        base / f"{prefix}_wt{root.wt_percent}_{job}.png",
    ]


def _table(doc: Document, title: str, headers: list[str], rows: list[list[str]]) -> None:
    p = doc.add_paragraph(title)
    if p.runs:
        p.runs[0].bold = True
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def _modulus_paired_rows(path: Path) -> list[list[str]]:
    if not path.exists():
        return []
    rows = _read_csv(path)
    out: list[list[str]] = []
    labels = {
        "gwy_stats": "Baseline Gwyddion stats",
        "gwy_ops_py_stats": "Gwyddion preprocess + Python stats",
        "raw_minmax": "Raw export + min/max filter",
        "raw_chauvenet": "Raw export + Chauvenet",
        "raw_three_sigma": "Raw export + 3-sigma",
    }
    for row in rows:
        out.append(
            [
                labels.get(row.get("method", ""), row.get("method", "")),
                row.get("n_pairs", ""),
                f"{_f(row.get('median_ratio_avg')):.3f}",
                f"{_f(row.get('mean_ratio_avg')):.3f}",
                f"{_f(row.get('median_delta_avg')) / 1e6:.3f}",
                f"{_f(row.get('mean_delta_avg')) / 1e6:.3f}",
                f"{_f(row.get('mean_delta_n_valid')):.1f}",
            ]
        )
    return out


def _poisson_success(mu: float, target: int) -> float:
    if mu <= 0:
        return 0.0
    term = math.exp(-mu)
    cdf = term
    for k in range(1, target):
        term *= mu / k
        cdf += term
    return max(0.0, min(1.0, 1.0 - cdf))


def _required_scans(lam: float, p: float, target: int, conf: float, max_scans: int) -> int:
    eff = lam * p
    if eff <= 0:
        return max_scans
    for n in range(1, max_scans + 1):
        if _poisson_success(eff * n, target) >= conf:
            return n
    return max_scans


def _availability_crossover_p(lam: float, available_scans: int, target: int, conf: float) -> float | None:
    if lam <= 0 or available_scans <= 0:
        return None
    if _poisson_success(lam * available_scans, target) < conf:
        return None
    low = 0.0
    high = 1.0
    for _ in range(60):
        mid = (low + high) / 2.0
        success = _poisson_success(lam * available_scans * mid, target)
        if success >= conf:
            high = mid
        else:
            low = mid
    return high


def _overview_plot(wt10: RootStats, wt25: RootStats) -> Path:
    import matplotlib.pyplot as plt

    out = ASSET_DIR / "chapter6_stage1_overview.png"
    out.parent.mkdir(parents=True, exist_ok=True)
    means_c = [wt10.method_stats[PRIMARY_JOB].mean_count, wt25.method_stats[PRIMARY_JOB].mean_count]
    stds_c = [wt10.method_stats[PRIMARY_JOB].std_count, wt25.method_stats[PRIMARY_JOB].std_count]
    means_i = [wt10.method_stats[PRIMARY_JOB].mean_isolated, wt25.method_stats[PRIMARY_JOB].mean_isolated]
    stds_i = [wt10.method_stats[PRIMARY_JOB].std_isolated, wt25.method_stats[PRIMARY_JOB].std_isolated]
    fig, ax = plt.subplots(figsize=(7.2, 4.6))
    x = [0, 1]
    w = 0.34
    ax.bar([a - w / 2 for a in x], means_c, w, yerr=stds_c, capsize=4, label="Candidate count/scan")
    ax.bar([a + w / 2 for a in x], means_i, w, yerr=stds_i, capsize=4, label="Isolated count/scan")
    ax.set_xticks(x, [wt10.label, wt25.label])
    ax.set_ylabel("Particles per scan")
    ax.set_title("Stage 1 overview\nPrimary-route means with standard deviations")
    ax.legend()
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(out, dpi=200)
    plt.close(fig)
    return out


def _crossover_plot(wt10: RootStats, wt25: RootStats) -> Path:
    import matplotlib.pyplot as plt

    out = ASSET_DIR / "chapter6_crossover_by_method.png"
    out.parent.mkdir(parents=True, exist_ok=True)
    ps = [i / 100 for i in range(1, 101)]
    fig, axes = plt.subplots(1, 2, figsize=(11.5, 4.8), sharey=True)
    for ax, root in zip(axes, [wt10, wt25]):
        available_scans = root.method_stats[PRIMARY_JOB].maps
        max_scans = max(available_scans, 220)
        label_offsets = {PRIMARY_JOB: -34, COMPARISON_JOB: 34}
        for job in JOB_ORDER:
            lam = root.method_fits[job].mean_per_scan
            ys = [_required_scans(lam, p, TARGET_ISOLATED, CONFIDENCE, max_scans) for p in ps]
            style = JOB_STYLES.get(job, {})
            ax.plot(
                ps,
                ys,
                linewidth=style.get("linewidth", 1.4),
                linestyle=style.get("linestyle", "-"),
                marker=style.get("marker"),
                markevery=max(1, int(len(ps) / 10.0)),
                markersize=4,
                color=style.get("color"),
                label=_job_label(job),
            )
            if job in (PRIMARY_JOB, COMPARISON_JOB):
                p_cross = _availability_crossover_p(lam, available_scans, TARGET_ISOLATED, CONFIDENCE)
                if p_cross is not None and 0.0 < p_cross <= 1.0:
                    y_cross = _required_scans(lam, p_cross, TARGET_ISOLATED, CONFIDENCE, max_scans)
                    highlight_marker = "*" if job == PRIMARY_JOB else style.get("marker", "o")
                    highlight_color = "#D4A017" if job == PRIMARY_JOB else style.get("color")
                    ax.scatter([p_cross], [y_cross], s=70 if job == PRIMARY_JOB else 40, zorder=6, color=highlight_color, edgecolor="black", linewidth=0.5, marker=highlight_marker)
                    ax.annotate(
                        f"{_job_label(job)}\np* = {p_cross:.3f}",
                        xy=(p_cross, y_cross),
                        xytext=(10, label_offsets.get(job, 10)),
                        textcoords="offset points",
                        fontsize=7,
                        va="center",
                        ha="left",
                        bbox=dict(boxstyle="round,pad=0.2", fc="white", ec=highlight_color, alpha=0.92),
                    )
        ax.axhline(available_scans, color="black", linestyle="--", linewidth=1.1, label="Available scans")
        ax.text(0.5, available_scans + 8, f"Available scans\n= {available_scans}", ha="center", va="bottom", fontsize=8, bbox=dict(boxstyle="round,pad=0.2", fc="white", ec="black", alpha=0.9))
        ax.set_title(f"{root.label}\nRequired scans vs confirmation probability")
        ax.set_xlabel("Confirmation probability p")
        ax.grid(alpha=0.25)
    axes[0].set_ylabel("Required scans for 95% confidence")
    handles, labels = axes[0].get_legend_handles_labels()
    fig.legend(handles, labels, loc="upper center", ncol=2, fontsize=8, frameon=True, facecolor="white", edgecolor="black")
    fig.tight_layout(rect=(0, 0, 1, 0.9))
    fig.savefig(out, dpi=200)
    plt.close(fig)
    return out


def _method_rows(root: RootStats) -> list[list[str]]:
    primary_mean = root.method_stats[PRIMARY_JOB].mean_isolated
    rows = []
    for job in JOB_ORDER:
        stats = root.method_stats[job]
        fit = root.method_fits[job]
        rows.append([
            root.label,
            _job_label(job),
            job,
            f"{stats.mean_isolated:.2f}",
            f"{stats.std_isolated:.2f}",
            _pct(stats.zero_isolated_rate),
            str(fit.n_required_095),
            f"{stats.mean_isolated / primary_mean:.2f}x",
            f"{stats.mean_isolated / root.wt_percent:.3f}",
        ])
    return rows


def _grain_rows(root: RootStats) -> list[list[str]]:
    rows = []
    for job in JOB_ORDER:
        grain = root.grain_stats[job]
        rows.append([
            root.label,
            _job_label(job),
            str(grain.grain_total),
            str(grain.grain_kept),
            str(grain.grain_isolated),
            _pm(grain.kept_mean_diameter_nm, grain.kept_std_diameter_nm, "nm"),
            _pm(grain.isolated_mean_diameter_nm, grain.isolated_std_diameter_nm, "nm"),
        ])
    return rows


def _grain_family_summary_rows(roots: list[RootStats], synthesis: dict | None = None) -> list[list[str]]:
    grouped: dict[tuple[str, str], list[dict[str, float]]] = defaultdict(list)

    def _family(job: str) -> str:
        if "medianbg" in job:
            return "Median background"
        if "flatten" in job:
            return "Flatten base"
        return "Other"

    source_rows = []
    if synthesis and synthesis.get("grain_rows"):
        for r in synthesis.get("grain_rows", []):
            source_rows.append({
                "system": r.get("system", ""),
                "job": r.get("job", ""),
                "grain_kept": float(r.get("grain_kept", 0) or 0),
                "grain_isolated": float(r.get("grain_isolated", 0) or 0),
                "kept_mean_diameter_nm": float(r.get("kept_mean_diameter_nm", 0) or 0),
                "kept_std_diameter_nm": float(r.get("kept_std_diameter_nm", 0) or 0),
                "isolated_mean_diameter_nm": float(r.get("isolated_mean_diameter_nm", 0) or 0),
                "isolated_std_diameter_nm": float(r.get("isolated_std_diameter_nm", 0) or 0),
            })
    else:
        for root in roots:
            for job in JOB_ORDER:
                grain = root.grain_stats[job]
                source_rows.append({
                    "system": root.label,
                    "job": job,
                    "grain_kept": float(grain.grain_kept),
                    "grain_isolated": float(grain.grain_isolated),
                    "kept_mean_diameter_nm": float(grain.kept_mean_diameter_nm),
                    "kept_std_diameter_nm": float(grain.kept_std_diameter_nm),
                    "isolated_mean_diameter_nm": float(grain.isolated_mean_diameter_nm),
                    "isolated_std_diameter_nm": float(grain.isolated_std_diameter_nm),
                })

    for row in source_rows:
        system = row["system"]
        family = _family(row["job"])
        grouped[(system, family)].append(row)
        grouped[(system, "All methods")].append(row)

    out_rows = []
    for system, family in sorted(grouped.keys()):
        rows = grouped[(system, family)]
        kept_counts = [r["grain_kept"] for r in rows]
        iso_counts = [r["grain_isolated"] for r in rows]
        kept_means = [r["kept_mean_diameter_nm"] for r in rows]
        kept_stds = [r["kept_std_diameter_nm"] for r in rows]
        iso_means = [r["isolated_mean_diameter_nm"] for r in rows]
        iso_stds = [r["isolated_std_diameter_nm"] for r in rows]
        out_rows.append([
            system,
            family,
            f"{mean(kept_counts):.1f}" if kept_counts else "",
            f"{mean(iso_counts):.1f}" if iso_counts else "",
            f"{mean(kept_means):.2f} +/- {mean(kept_stds):.2f}" if kept_means and kept_stds else "",
            f"{mean(iso_means):.2f} +/- {mean(iso_stds):.2f}" if iso_means and iso_stds else "",
        ])
    return out_rows


def _family_scan_requirement_rows(system_label: str, roots: list[RootStats], synthesis: dict | None = None) -> list[list[str]]:
    def _family(job: str) -> str:
        if "medianbg" in job:
            return "Median background"
        if "flatten" in job:
            return "Flatten base"
        return "Other"

    source_rows = []
    if synthesis and synthesis.get("method_rows"):
        for r in synthesis.get("method_rows", []):
            if r.get("system", "") != system_label:
                continue
            source_rows.append({
                "system": r.get("system", ""),
                "method": r.get("method", ""),
                "job": r.get("job", ""),
                "mean_isolated_per_scan": float(r.get("mean_isolated_per_scan", 0) or 0),
                "std_isolated_per_scan": float(r.get("std_isolated_per_scan", 0) or 0),
                "required_scans_095": str(r.get("required_scans_095", "")),
            })
    else:
        for root in roots:
            if root.label != system_label:
                continue
            for job in JOB_ORDER:
                stats = root.method_stats[job]
                fit = root.method_fits[job]
                source_rows.append({
                    "system": root.label,
                    "method": _job_label(job),
                    "job": job,
                    "mean_isolated_per_scan": stats.mean_isolated,
                    "std_isolated_per_scan": stats.std_isolated,
                    "required_scans_095": str(fit.n_required_095),
                })

    rows = []
    for row in source_rows:
        rows.append([
            _family(row["job"]),
            row["method"],
            f"{row['mean_isolated_per_scan']:.2f}",
            f"{row['std_isolated_per_scan']:.2f}",
            row["required_scans_095"],
        ])
    return rows


def _sample_rows(root: RootStats) -> list[list[str]]:
    stats = root.method_stats[PRIMARY_JOB]
    rows = []
    for sample in root.samples:
        rows.append([
            root.label,
            sample.replace("_", " "),
            str(stats.sample_scan_counts[sample]),
            f"{stats.sample_means[sample]:.2f}",
            f"{stats.sample_stds[sample]:.2f}",
            f"{stats.sample_pct_nonzero[sample]:.1f}%",
        ])
    return rows


def _write(doc_path: Path, wt10: RootStats, wt25: RootStats) -> None:
    synthesis_path = wt10.root.parent / "topo_report_synthesis.json"
    synthesis = None
    if synthesis_path.exists():
        try:
            synthesis = json.loads(synthesis_path.read_text(encoding="utf-8"))
        except Exception:
            synthesis = None
    if synthesis is None:
        try:
            synthesis = build_bundle(wt10.root, wt25.root)
        except Exception:
            synthesis = None

    overview = _overview_plot(wt10, wt25)
    crossover = _crossover_plot(wt10, wt25)
    p10 = wt10.method_stats[PRIMARY_JOB]
    p25 = wt25.method_stats[PRIMARY_JOB]
    f10 = wt10.method_fits[PRIMARY_JOB]
    f25 = wt25.method_fits[PRIMARY_JOB]
    c10 = wt10.method_stats[COMPARISON_JOB]
    c25 = wt25.method_stats[COMPARISON_JOB]
    cf10 = wt10.method_fits[COMPARISON_JOB]
    cf25 = wt25.method_fits[COMPARISON_JOB]

    doc = Document()
    doc.add_heading("Chapter 6 Draft - Stage 1 Results and Feasibility Decision", level=0)
    doc.add_paragraph(f"Draft generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("Scope: measurement feasibility, scan efficiency, and decision thresholds only.")
    if synthesis is not None:
        doc.add_paragraph(f"Shared synthesis source: {synthesis_path if synthesis_path.exists() else 'in-memory synthesis bundle'}")

    doc.add_heading("6.1 Processing Validation - Baseline PEGDA (No SiNP)", level=1)
    doc.add_paragraph("Baseline PEGDA validation in this chapter is carried primarily by the modulus method-comparison study. That study was used to evaluate forward/backward agreement, route consistency, and the practical effect of pixel loss before extending the Stage 1 logic to the particle-count workflow.")
    doc.add_paragraph("The modulus baseline artifact is limited in scope: it was performed on one PEGDA sample and one fracture-surface side, so it should be interpreted as method-validation evidence rather than a population-level PEGDA modulus study. Even with that scope limit, the paired forward/backward comparison and the route-consistency summaries provide the direct baseline-validation evidence requested in the thesis outline.")
    doc.add_paragraph("Unit provenance should be stated explicitly for this section. In this workflow, Gwyddion-derived outputs inherit the physical unit attached to the active data field when that unit is actually present in the imported field metadata. A direct one-file verification run on the current PEGDA modulus TIFFs showed that pygwy did not detect an embedded modulus z-unit for that file, so the current modulus source summaries carrying units = kPa should be treated as a workflow-level fallback/default assignment rather than fully verified source metadata truth. That same one-file verification run also produced a negative modulus export, so the absolute modulus value path should be treated as provisional pending targeted re-validation of the TIFF unit/processing chain. By contrast, the Stage 1 particle-count outputs are count-based summaries and therefore do not inherit a physical z-unit, while particle and grain diameter summaries are derived geometric metrics reported in nm.")
    _table(
        doc,
        "Table 6.1 - Paired forward/backward modulus baseline summary by method",
        ["Method", "Paired scans", "Median B/F ratio", "Mean B/F ratio", "Median delta avg (GPa-equiv.)", "Mean delta avg (GPa-equiv.)", "Mean delta n_valid"],
        _modulus_paired_rows(MODULUS_PAIRED_SUMMARY),
    )
    _figure(
        doc,
        "Figure 6.1 - Baseline PEGDA modulus validation figures. The paired-ratio box plot summarizes the combined forward/backward comparison by method, while the absolute-modulus bars show the forward-only and backward-only method summaries with standard-error bars. Together these figures support route consistency and the forward-only selection used later in the Stage 1 topography workflow.",
        [
            MODULUS_ASSET_DIR / "paired_ratio_boxplot.png",
            MODULUS_ASSET_DIR / "forward_absolute_modulus_bar.png",
            MODULUS_ASSET_DIR / "backward_absolute_modulus_bar.png",
        ],
        width=5.5,
    )
    _figure(
        doc,
        "Figure 6.1b - Forward baseline, backward baseline, and combined forward/backward PEGDA modulus heatmaps plus one representative forward delta-versus-baseline comparison heatmap. These plots show the spatial baseline field for each scan direction, the paired mean baseline field, and one example of how a comparison route shifts the forward baseline modulus pattern.",
        [
            MODULUS_FORWARD_COMPARE_DIR / "plots" / "heatmap_two_panel__baseline.png",
            MODULUS_BACKWARD_COMPARE_DIR / "plots" / "heatmap_two_panel__baseline.png",
            MODULUS_ASSET_DIR / "combined_forward_backward_baseline_heatmap.png",
            MODULUS_FORWARD_COMPARE_DIR / "plots" / "heatmap_two_panel__delta_vs_baseline__raw_minmax.png",
        ],
        width=5.5,
    )
    doc.add_heading("6.2 Stage 1 - Particle Presence in PEGDA-SiNP", level=1)
    doc.add_heading("6.2.1 Scan Inventory", level=2)
    _table(doc, "Table 6.2 - Scan inventory used for the current Stage 1 analysis", ["System", "Sample sets", "Scans analyzed", "Scan size (um x um)", "Pixel grid", "Resolution (nm/pixel)", "Nominal map grid"], [[wt10.label, str(len(wt10.samples)), str(p10.maps), SCAN_SIZE_UM, PIXEL_GRID, RESOLUTION_NM_PER_PIXEL, NOMINAL_GRID], [wt25.label, str(len(wt25.samples)), str(p25.maps), SCAN_SIZE_UM, PIXEL_GRID, RESOLUTION_NM_PER_PIXEL, NOMINAL_GRID]])
    doc.add_paragraph(f"{wt10.label} contributed {p10.maps} scans across {len(wt10.samples)} sample sets. {wt25.label} contributed {p25.maps} scans across {len(wt25.samples)} sample sets. Each scan is a 5 um x 5 um AFM image on a nominal 21 x 21 survey grid with 512 x 512 pixels.")
    doc.add_paragraph("The nominal survey layout was 21 x 21 scans, but the analyzed inventories reflect only scans that were actually collected and retained in the grouped dataset. Some 25 wt% grouped sample sets were therefore incomplete relative to the nominal grid, even though the total analyzed inventory remained well above the later scan-sufficiency threshold.")

    doc.add_heading("6.2.2 Particle Count Per Scan", level=2)
    doc.add_paragraph(f"Under the primary route ({PRIMARY_JOB}), retained candidate counts averaged {_pm(p10.mean_count, p10.std_count)} particles per scan for {wt10.label} and {_pm(p25.mean_count, p25.std_count)} particles per scan for {wt25.label}. Zero-count scans were uncommon in both grouped datasets ({_pct(p10.zero_count_rate)} and {_pct(p25.zero_count_rate)}).")
    doc.add_paragraph(f"The {wt25.label} dataset showed a modest increase in candidate-particle density relative to {wt10.label} ({p25.mean_count / p10.mean_count:.2f}x), but this stage still represents Stage 1 candidate features rather than confirmed particles.")
    doc.add_paragraph("The count histograms show that low candidate counts remain the most common outcome for both loadings, indicating that the retained population is typically distributed as small numbers of particle candidates per scan rather than as large multi-particle fields.")
    _table(doc, "Table 6.3 - Primary-method candidate count summary", ["System", "Mean count/scan", "Std", "Min", "Max", "Zero-count rate"], [[wt10.label, f"{p10.mean_count:.2f}", f"{p10.std_count:.2f}", str(p10.min_count), str(p10.max_count), _pct(p10.zero_count_rate)], [wt25.label, f"{p25.mean_count:.2f}", f"{p25.std_count:.2f}", str(p25.min_count), str(p25.max_count), _pct(p25.zero_count_rate)]])
    _figure(doc, "Figure 6.2a - Stage 1 overview for PEGDA, 1 wt% TPO, no coating, comparing 10 wt% and 25 wt% SiNP. Bars show mean candidate and isolated counts per scan with standard deviations.", [overview])
    _figure(doc, "Figure 6.2b - Particle-count histograms for PEGDA, 1 wt% TPO, no coating, with 10 wt% and 25 wt% SiNP under the primary route.", [wt10.root / "fig_particle_count_hist.png", wt25.root / "fig_particle_count_hist.png"])
    _figure(
        doc,
        "Figure 6.2c - Representative raw topography image showing a mixed particle field before the mean kept-particle density maps are introduced. This figure provides visual context for the kinds of particle arrangements counted during Stage 1.",
        [REP_IMAGE_DIR / "pegda10_all_three_types.png"],
        width=5.2,
    )
    _figure(doc, "Figure 6.3 - Particle-count grid heatmaps for the primary route, separated by SiNP loading.", [wt10.root / "summary_outputs" / "combined" / f"fig_particle_count_grid_wt10_{PRIMARY_JOB}.png", wt25.root / "summary_outputs" / "combined" / f"fig_particle_count_grid_wt25_{PRIMARY_JOB}.png"])
    _figure(
        doc,
        "Figure 6.3b - Standard-deviation companions for the mean kept-particle heatmaps under the primary route. These maps show the between-sample spread at each scan position for the 10 wt% and 25 wt% grouped datasets.",
        [
            wt10.root / "summary_outputs" / "combined" / f"fig_particle_count_grid_std_wt10_{PRIMARY_JOB}.png",
            wt25.root / "summary_outputs" / "combined" / f"fig_particle_count_grid_std_wt25_{PRIMARY_JOB}.png",
        ],
        width=5.5,
    )

    doc.add_heading("6.2.3 Particle Diameter Distribution", level=2)
    doc.add_paragraph(f"After applying the configured {DIAMETER_FILTER_NM} nm filter, retained particle diameter averaged {_pm(wt10.diameter_mean_nm, wt10.diameter_std_nm, 'nm')} for {wt10.label} and {_pm(wt25.diameter_mean_nm, wt25.diameter_std_nm, 'nm')} for {wt25.label}. The retained means remain centered on the expected SiNP size band.")
    _table(doc, "Table 6.4 - Retained particle diameter summary", ["System", "Filter band (nm)", "Mean diameter (nm)", "Std diameter (nm)"], [[wt10.label, DIAMETER_FILTER_NM, f"{wt10.diameter_mean_nm:.2f}", f"{wt10.diameter_std_nm:.2f}"], [wt25.label, DIAMETER_FILTER_NM, f"{wt25.diameter_mean_nm:.2f}", f"{wt25.diameter_std_nm:.2f}"]])
    _figure(doc, "Figure 6.4 - Retained particle-diameter histograms for PEGDA, 1 wt% TPO, no coating, comparing 10 wt% and 25 wt% SiNP after filtering.", [wt10.root / "fig_particle_diameter_hist.png", wt25.root / "fig_particle_diameter_hist.png"])

    doc.add_heading("6.3 Isolation Analysis", level=1)
    doc.add_heading("6.3.1 Isolation Count Per Scan", level=2)
    doc.add_paragraph(f"Isolation is the controlling Stage 1 feasibility metric because only spatially separated targets are usable for Stage 2. In the current workflow, isolation is defined by a minimum center-to-center spacing of {ISOLATION_DISTANCE_NM} nm.")
    doc.add_paragraph(f"Primary-route isolated counts averaged {_pm(p10.mean_isolated, p10.std_isolated)} per scan for {wt10.label} and {_pm(p25.mean_isolated, p25.std_isolated)} per scan for {wt25.label}. Scans with at least one isolated particle accounted for {p10.pct_with_isolated:.1f}% and {p25.pct_with_isolated:.1f}% of the two datasets, respectively.")
    doc.add_paragraph(f"The higher-loading dataset did not produce a proportionally larger isolated-particle yield. Mean isolated counts remained close ({p10.mean_isolated:.2f} versus {p25.mean_isolated:.2f} isolated candidates per scan) even though candidate density increased.")
    doc.add_paragraph("The isolated-count maps summarize mean isolated-candidate count at each scan position across the grouped sample sets. They are therefore useful for showing spatial pattern, while the paired standard-deviation maps show how much that isolated-candidate density varies between grouped samples at the same position.")
    sample_table_rows = _sample_rows(wt10) + _sample_rows(wt25)
    if synthesis and synthesis.get("sample_rows"):
        sample_table_rows = [[
            r.get("system", ""),
            r.get("sample_set", "").replace("_", " "),
            r.get("scans", ""),
            r.get("mean_isolated_per_scan", ""),
            r.get("std_isolated_per_scan", ""),
            f"{r.get('pct_scans_with_isolated', '')}%",
        ] for r in synthesis.get("sample_rows", [])]
    _table(doc, "Table 6.5 - Sample-level isolated-particle yield under the primary route", ["System", "Sample set", "Scans", "Mean isolated/scan", "Std", "% scans with >=1 isolated"], sample_table_rows)
    _figure(doc, "Figure 6.5 - Histograms of isolated-particle counts per scan for the primary route, separated by 10 wt% and 25 wt% SiNP.", [wt10.root / "fig_isolated_count_hist.png", wt25.root / "fig_isolated_count_hist.png"])
    _figure(doc, "Figure 6.6 - Isolated-particle grid heatmaps for the primary route, separated by SiNP loading.", [wt10.root / "summary_outputs" / "combined" / f"fig_isolated_count_grid_wt10_{PRIMARY_JOB}.png", wt25.root / "summary_outputs" / "combined" / f"fig_isolated_count_grid_wt25_{PRIMARY_JOB}.png"])
    _figure(
        doc,
        "Figure 6.6b - Standard-deviation companions for the mean isolated-particle heatmaps under the primary route. These maps show the between-sample spread in isolated-candidate density at each scan position for the 10 wt% and 25 wt% grouped datasets.",
        [
            wt10.root / "summary_outputs" / "combined" / f"fig_isolated_count_grid_std_wt10_{PRIMARY_JOB}.png",
            wt25.root / "summary_outputs" / "combined" / f"fig_isolated_count_grid_std_wt25_{PRIMARY_JOB}.png",
        ],
        width=5.5,
    )

    doc.add_heading("6.3.2 Required Scans for 95% Confidence", level=2)
    doc.add_paragraph(f"Using the Poisson baseline fit, a target of {TARGET_ISOLATED} isolated particles, and {int(CONFIDENCE * 100)}% confidence, the primary route required {f10.n_required_095} scans for {wt10.label} and {f25.n_required_095} scans for {wt25.label}. Both analyzed inventories greatly exceed those requirements.")
    doc.add_paragraph("In Table 6.5, primary lambda is the mean isolated-particle count per scan under the primary route. The observed zero-isolated rate is the empirical fraction of scans with zero isolated particles. The required-scan value is the smallest number of scans for which the modeled probability of obtaining at least 30 isolated particles reaches 95%.")
    required_rows = [[wt10.label, f"{f10.mean_per_scan:.3f}", _pct(f10.zero_rate_obs), str(f10.n_required_095), str(p10.maps)], [wt25.label, f"{f25.mean_per_scan:.3f}", _pct(f25.zero_rate_obs), str(f25.n_required_095), str(p25.maps)]]
    if synthesis and synthesis.get("required_scan_rows"):
        required_rows = [[
            r.get("system", ""),
            r.get("primary_lambda", ""),
            _pct(float(r.get("observed_zero_isolated_rate", 0) or 0)),
            r.get("required_scans_095", ""),
            r.get("available_scans", ""),
        ] for r in synthesis.get("required_scan_rows", [])]
    _table(doc, "Table 6.6 - Required scans for isolated-particle sufficiency under the primary route", ["System", "Primary lambda", "Observed zero-isolated rate", "Required scans (95%)", "Available scans"], required_rows)

    doc.add_heading("6.4 Grain Metrics", level=1)
    doc.add_paragraph("Grain exports are available for the full method matrix. These results are used here as supporting evidence for segmentation consistency across methods; they do not replace the particle-count or isolation metrics that control the Stage 1 feasibility decision.")
    doc.add_paragraph("In this section, the grain summary table and the grain diameter distribution plots are intended to show whether the extracted grain populations remain broadly comparable across methods. The purpose is not to define feasibility from grain statistics alone, but to show that the underlying segmentation behavior is not changing so dramatically that the particle-based conclusions become uninterpretable.")
    grain_table_rows = _grain_rows(wt10) + _grain_rows(wt25)
    if synthesis and synthesis.get("grain_rows"):
        grain_table_rows = [[
            r.get("system", ""),
            r.get("method", ""),
            r.get("grain_total", ""),
            r.get("grain_kept", ""),
            r.get("grain_isolated", ""),
            f"{r.get('kept_mean_diameter_nm', '')} +/- {r.get('kept_std_diameter_nm', '')} nm",
            f"{r.get('isolated_mean_diameter_nm', '')} +/- {r.get('isolated_std_diameter_nm', '')} nm",
        ] for r in synthesis.get("grain_rows", [])]
    _table(doc, "Table 6.7 - Grain summary across the full method matrix", ["System", "Method", "Grain total", "Grain kept", "Grain isolated", "Kept diameter mean +/- std", "Isolated diameter mean +/- std"], grain_table_rows)
    family_grain_rows = _grain_family_summary_rows([wt10, wt25], synthesis)
    _table(doc, "Table 6.7b - Grain summary by processing family", ["System", "Processing family", "Avg. grain kept", "Avg. grain isolated", "Avg. kept diameter mean +/- std", "Avg. isolated diameter mean +/- std"], family_grain_rows)
    _figure(doc, "Figure 6.7 - Grain diameter distributions by method for PEGDA, 1 wt% TPO, no coating, separated by 10 wt% and 25 wt% SiNP. These summaries are used as supporting evidence that the segmented grain populations remain broadly comparable across methods.", [wt10.root / "summary_outputs" / "grain_compare" / "fig_grain_diameter_nm_kept_box_by_job.png", wt25.root / "summary_outputs" / "grain_compare" / "fig_grain_diameter_nm_kept_box_by_job.png"], width=5.7)
    _figure(doc, "Figure 6.7b - Grain diameter summary bars by method for PEGDA, 1 wt% TPO, no coating, separated by 10 wt% and 25 wt% SiNP. These plots show method-level mean and standard deviation, while the box plots show the corresponding distribution spread.", [wt10.root / "summary_outputs" / "grain_compare" / "fig_grain_diameter_nm_kept_mean_by_job.png", wt25.root / "summary_outputs" / "grain_compare" / "fig_grain_diameter_nm_kept_mean_by_job.png"], width=5.7)

    doc.add_heading("6.5 Processing Route Sensitivity", level=1)
    doc.add_paragraph(f"The practical sensitivity question is not whether the processing routes produce visually different maps, but whether they materially change isolated-particle yield and therefore the number of scans needed for Stage 1 sufficiency. Relative to the primary route, the comparison route ({COMPARISON_JOB}) reduced mean isolated yield from {_pm(p10.mean_isolated, p10.std_isolated)} to {_pm(c10.mean_isolated, c10.std_isolated)} per scan in {wt10.label} and from {_pm(p25.mean_isolated, p25.std_isolated)} to {_pm(c25.mean_isolated, c25.std_isolated)} per scan in {wt25.label}. The corresponding 95% scan requirements increased from {f10.n_required_095} to {cf10.n_required_095} scans and from {f25.n_required_095} to {cf25.n_required_095} scans.")
    doc.add_paragraph("The methods fall into two processing families: a median-background family and a flatten-base family. Tables 6.8a and 6.8b summarize those families within each loading by listing the submethod, the isolated-particle yield, and the number of scans required to reach 30 isolated Stage 1 particle candidates at 95% modeled confidence. These compact tables provide the direct numerical counterpart to the family-comparison plots.")
    doc.add_paragraph("The full method table should be interpreted provisionally. A targeted post-generation verification identified a mask write-back issue in the thresholded particle-analysis path, so the full all-method matrix must be regenerated after that fix before any final claim is made about whether threshold variants truly collapse within a preprocessing family.")
    _table(
        doc,
        "Table 6.8a - 10 wt% processing-family scan requirements for 30 isolated candidates at 95% confidence",
        ["Processing family", "Submethod", "Mean isolated/scan", "Std", "Req. scans (95%)"],
        _family_scan_requirement_rows(wt10.label, [wt10, wt25], synthesis),
    )
    _table(
        doc,
        "Table 6.8b - 25 wt% processing-family scan requirements for 30 isolated candidates at 95% confidence",
        ["Processing family", "Submethod", "Mean isolated/scan", "Std", "Req. scans (95%)"],
        _family_scan_requirement_rows(wt25.label, [wt10, wt25], synthesis),
    )
    method_table_rows = _method_rows(wt10) + _method_rows(wt25)
    if synthesis and synthesis.get("method_rows"):
        method_table_rows = [[
            r.get("system", ""),
            r.get("method", ""),
            r.get("job", ""),
            r.get("mean_isolated_per_scan", ""),
            r.get("std_isolated_per_scan", ""),
            _pct(float(r.get("zero_isolated_rate", 0) or 0)),
            r.get("required_scans_095", ""),
            r.get("relative_to_primary", ""),
            r.get("isolated_per_scan_per_wt", ""),
        ] for r in synthesis.get("method_rows", [])]
    _table(doc, "Table 6.8 - All-method comparison with loading-normalized isolation yield", ["System", "Method", "Job/profile", "Mean isolated/scan", "Std", "Zero-isolated rate", "Req. scans (95%)", "Relative to primary", "Isolated/scan/wt%"], method_table_rows)
    _figure(doc, "Figure 6.8 - Full-method comparison of kept and isolated counts by processing family and SiNP loading for PEGDA, 1 wt% TPO, no coating. Separate median-background and flatten-base panels are shown for each wt% dataset so that submethod differences are visible without complete overlap.", [wt10.root / "summary_outputs" / "compare_by_wt" / "fig_method_family_counts_10pct.png", wt25.root / "summary_outputs" / "compare_by_wt" / "fig_method_family_counts_25pct.png"])
    _figure(
        doc,
        "Figure 6.9 - Aggregate Poisson success-probability curves for all methods, separated by SiNP loading and scrape state. Horizontal dashed lines mark the 90%, 95%, and 99% success thresholds; vertical markers show the 95%-requirement scan count for each method.",
        [
            wt10.root / "summary_outputs" / "fits" / "risk_aggregate_10____Non-scraped_poisson.png",
            wt10.root / "summary_outputs" / "fits" / "risk_aggregate_10____Scraped_poisson.png",
            wt25.root / "summary_outputs" / "fits" / "risk_aggregate_25____Non-scraped_poisson.png",
            wt25.root / "summary_outputs" / "fits" / "risk_aggregate_25____Scraped_poisson.png",
        ],
        width=5.2,
    )

    doc.add_heading("6.6 Stage 2 Trigger / Crossover Decision", level=1)
    p_cross_10 = _availability_crossover_p(f10.mean_per_scan, p10.maps, TARGET_ISOLATED, CONFIDENCE)
    p_cross_25 = _availability_crossover_p(f25.mean_per_scan, p25.maps, TARGET_ISOLATED, CONFIDENCE)
    doc.add_paragraph("The Stage 2 trigger is framed as a sensitivity study in the confirmation probability p. The practical question is how far the confirmation rate could fall before the current Stage 1 scan inventory would no longer meet the 30-particle, 95%-confidence target.")
    doc.add_paragraph("In Figure 6.10, the x-axis is the assumed fraction of Stage 1 isolated candidates that would later be confirmed as true particles in Stage 2, and the y-axis is the scan count required to still meet the Stage 2 target. The horizontal reference line is the currently available scan inventory. The crossover value p* is therefore the minimum confirmation fraction required for the existing inventory to remain sufficient; it is not the probability of a crossover event.")
    crossover_rows = [[wt10.label, f"{f10.mean_per_scan:.3f}", str(p10.maps), f"{p_cross_10:.3f}" if p_cross_10 is not None else "not reachable"], [wt25.label, f"{f25.mean_per_scan:.3f}", str(p25.maps), f"{p_cross_25:.3f}" if p_cross_25 is not None else "not reachable"]]
    if synthesis and synthesis.get("crossover_rows"):
        crossover_rows = [[
            r.get("system", ""),
            r.get("primary_lambda", ""),
            r.get("available_scans", ""),
            r.get("minimum_confirmation_fraction_pstar", ""),
        ] for r in synthesis.get("crossover_rows", [])]
    _table(doc, "Table 6.9 - Minimum confirmation fraction p* required for the current scan inventory to remain sufficient", ["System", "Primary lambda", "Available scans", "Minimum confirmation fraction p*"], crossover_rows)
    _figure(doc, "Figure 6.10a - Required scan count versus confirmation probability p across the full method matrix, with the available-scan inventory shown as a horizontal reference and labeled crossover points for the primary and comparison routes. PEGDA, 1 wt% TPO, no coating, separated by 10 wt% and 25 wt% SiNP.", [crossover], width=6.6)
    _figure(
        doc,
        "Figure 6.10b - Representative Stage 2 context images from the 0 wt% SiNP contact-mode dataset. These images provide visual context for the multi-channel interrogation workflow that would be applied once Stage 1 isolated targets are selected.",
        [
            REP_IMAGE_DIR / "stage2_topography_channel.png",
            REP_IMAGE_DIR / "stage2_modulus_channel.png",
            REP_IMAGE_DIR / "stage2_adhesion_channel.png",
            REP_IMAGE_DIR / "stage2_deformation_channel.png",
        ],
        width=5.2,
    )

    doc.add_heading("6.7 Stage 1 Decision", level=1)
    doc.add_paragraph(f"For {wt10.label}, candidate particle presence was confirmed, the retained diameter distribution remained consistent with the configured {DIAMETER_FILTER_NM} nm size window, isolated particles occurred in {p10.pct_with_isolated:.1f}% of scans under the primary route, and the analyzed inventory of {p10.maps} scans greatly exceeded the {f10.n_required_095}-scan requirement.")
    doc.add_paragraph(f"For {wt25.label}, candidate particle presence was likewise confirmed, the retained diameter distribution remained acceptable, isolated particles occurred in {p25.pct_with_isolated:.1f}% of scans under the primary route, and the analyzed inventory of {p25.maps} scans greatly exceeded the corresponding {f25.n_required_095}-scan requirement.")
    doc.add_paragraph("Stage 2 high-resolution interrogation is therefore justified for both loadings under the current Stage 1 interpretation. The comparative result is that higher loading increased candidate density modestly, but did not materially improve isolated-particle availability enough to change the decision.")

    doc.add_heading("6.8 Discussion and Limits", level=1)
    doc.add_paragraph("Within the current Stage 1 scope, two conclusions are supported directly by the data. First, isolated candidates are available in sufficient number under the primary route. Second, increasing SiNP loading from 10 wt% to 25 wt% does not automatically yield a proportionally larger population of isolated targets. The remaining uncertainty concerns how many of those isolated candidates would later be confirmed as true particles during Stage 2 multi-channel interrogation.")
    doc.add_paragraph("The crossover analysis is therefore useful because it translates that remaining uncertainty into an explicit scan-budget threshold rather than leaving Stage 2 as an undefined next step. In this sense, Stage 2 remains a neutral unknown to be tested rather than an assumed confirmation of every Stage 1 candidate.")
    doc.add_paragraph("A second limitation concerns model justification. The present inventories are comfortably above the scan counts required for the immediate Stage 1 feasibility claim, even though the two loadings do not contain equal numbers of scans. That supports the practical sufficiency argument used here, but it does not by itself establish that Poisson is the only admissible count model for later extensions of the workflow.")

    doc_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(doc_path)


def main() -> int:
    parser = argparse.ArgumentParser(description="Populate the Chapter 6 thesis draft from current Stage 1 output roots.")
    parser.add_argument("--wt10-root", required=True)
    parser.add_argument("--wt25-root", required=True)
    parser.add_argument("--docx-path", default="docs/Thesis/Chapter6_Stage1_Results_Feasibility_DRAFT.docx")
    args = parser.parse_args()
    _write(Path(args.docx_path), _load_root(Path(args.wt10_root), 10), _load_root(Path(args.wt25_root), 25))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

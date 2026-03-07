from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches
from PIL import Image


BASE = Path(r"C:\Users\Conor O'Brien\Dropbox\03_AML\00 IN-BOX\AFM Topo Particle processing OUT\Particle Images for Thesis Figures")
RAW = BASE / "raw"
PROCESSED = BASE / "processed"

CH4 = Path("docs/Thesis/Chapter4_AFM_Acquisition_Protocol_DRAFT.docx")
CH5 = Path("docs/Thesis/Chapter5_Statistical_Validation_Framework_DRAFT_v1alpha-20260306-1854.docx")
CH6 = Path("docs/Thesis/Chapter6_Stage1_Results_Feasibility_DRAFT_v1alpha-20260306-2114.docx")


IMAGES = {
    "probe_vision": {
        "raw_name": "PEGDA01TPO10SiNPSam01_S2_P__GrID018_LOC_RC004014--Vision_-250929-CRO.jpg",
        "out_name": "probe_vision.png",
        "chapter": 4,
        "heading": "Representative Acquisition Support Images",
        "title": "Probe vision image for fractured PEGDA with 10 wt% SiNP.",
        "caption": (
            "Representative probe-vision image associated with fractured PEGDA containing 10 wt% SiNP. "
            "This figure is included as acquisition-context support material rather than a quantitative analysis image."
        ),
    },
    "calibration_lithography": {
        "raw_name": "HS100MG-PPP-NCM_250925_ChNameDirVision__ScanSize_GrID001_LocationOfScanRC-CRO.jpg",
        "out_name": "calibration_lithography.png",
        "chapter": 4,
        "heading": "Representative Acquisition Support Images",
        "title": "Silicon lithography calibration image.",
        "caption": (
            "Silicon lithography image used for calibration support. This figure is useful for documenting the "
            "instrument/calibration context of the AFM workflow."
        ),
    },
    "sinp_only_topo": {
        "raw_name": "SiNP-Calibration_GrID011_LOC_RC003003-5.00x5.00-Z Height_Forward-250925-CRO.tiff",
        "out_name": "sinp_only_topography.png",
        "chapter": 5,
        "heading": "Representative Topography Image Examples",
        "title": "SiNP-only topography example.",
        "caption": (
            "Representative SiNP-only topography image. The raw image shows scarring, so this figure should be treated "
            "as a preprocessing/editing example rather than a finalized quantitative reference."
        ),
    },
    "pegda10_all_three": {
        "raw_name": "PEGDA01TPO025SiNP_Sam01_S2_P__GrID007_LOC_RC001016-5.00x5.00-Z Height_Forward-251008-CRO.tiff",
        "out_name": "pegda10_all_three_types.png",
        "chapter": 5,
        "heading": "Representative Topography Image Examples",
        "title": "PEGDA-SiNP topography example containing multiple particle arrangements.",
        "caption": (
            "Representative PEGDA-SiNP topography image noted to contain all three major particle-field types of interest. "
            "This figure is intended to help explain what the Stage 1 workflow is detecting visually."
        ),
    },
    "stage2_topo": {
        "raw_name": "PEGDA01TPO000SiNP_Sam01_S1_P__GrID079_LOC_RC016021-5.00x5.00-Z Height_Forward-251102-CRO.tiff",
        "out_name": "stage2_topography_channel.png",
        "chapter": 6,
        "heading": "Representative Multi-Channel Stage 2 Context Images",
        "title": "Representative contact-mode topography channel.",
        "caption": (
            "Representative 0 wt% SiNP contact-mode topography channel. This image is included as Stage 2 context for "
            "how multi-channel information could later be used to support mask construction or particle confirmation."
        ),
    },
    "stage2_modulus": {
        "raw_name": "PEGDA01TPO000SiNP_Sam01_S1_P__GrID079_LOC_RC016021-5.00x5.00-Modulus_Forward-251102-CRO.tiff",
        "out_name": "stage2_modulus_channel.png",
        "chapter": 6,
        "heading": "Representative Multi-Channel Stage 2 Context Images",
        "title": "Representative contact-mode modulus channel.",
        "caption": (
            "Representative 0 wt% SiNP contact-mode modulus channel shown as Stage 2 context only. "
            "This figure illustrates one of the additional channels that could support future multi-channel validation."
        ),
    },
    "stage2_adhesion": {
        "raw_name": "PEGDA01TPO000SiNP_Sam01_S1_P__GrID079_LOC_RC016021-5.00x5.00-Adhesion Energy_Forward-251102-CRO.tiff",
        "out_name": "stage2_adhesion_channel.png",
        "chapter": 6,
        "heading": "Representative Multi-Channel Stage 2 Context Images",
        "title": "Representative contact-mode adhesion-energy channel.",
        "caption": (
            "Representative 0 wt% SiNP contact-mode adhesion-energy channel shown as Stage 2 context. "
            "This figure is useful for illustrating how additional AFM channels may contribute to future mask design."
        ),
    },
    "stage2_deformation": {
        "raw_name": "PEGDA01TPO000SiNP_Sam01_S1_P__GrID079_LOC_RC016021-5.00x5.00-Deformation_Forward-251102-CRO.tiff",
        "out_name": "stage2_deformation_channel.png",
        "chapter": 6,
        "heading": "Representative Multi-Channel Stage 2 Context Images",
        "title": "Representative contact-mode deformation channel.",
        "caption": (
            "Representative 0 wt% SiNP contact-mode deformation channel shown as Stage 2 context. "
            "Together with the other contact-mode channels, this image helps illustrate the future multi-channel validation concept."
        ),
    },
}


def convert_for_docx(src: Path, dst: Path) -> Path:
    dst.parent.mkdir(parents=True, exist_ok=True)
    with Image.open(src) as im:
        if getattr(im, "n_frames", 1) > 1:
            im.seek(0)
        rgb = im.convert("RGB")
        rgb.save(dst, format="PNG")
    return dst


def add_images(doc_path: Path, chapter_num: int, heading: str, image_keys: list[str], out_name: str) -> None:
    doc = Document(doc_path)
    doc.add_page_break()
    doc.add_heading(heading, level=1)
    for key in image_keys:
        item = IMAGES[key]
        src = RAW / item["raw_name"]
        dst = PROCESSED / item["out_name"]
        convert_for_docx(src, dst)
        doc.add_paragraph(item["title"])
        doc.add_picture(str(dst), width=Inches(5.8))
        doc.add_paragraph(item["caption"])
        doc.add_paragraph(f"Source: {src}")
    target = doc_path.with_name(out_name)
    doc.save(target)


def main() -> int:
    add_images(
        CH4,
        4,
        "Representative Acquisition Support Images",
        ["probe_vision", "calibration_lithography"],
        "Chapter4_AFM_Acquisition_Protocol_DRAFT_v1alpha-20260306-figureimport.docx",
    )
    add_images(
        CH5,
        5,
        "Representative Topography Image Examples",
        ["sinp_only_topo", "pegda10_all_three"],
        "Chapter5_Statistical_Validation_Framework_DRAFT_v1alpha-20260306-figureimport.docx",
    )
    add_images(
        CH6,
        6,
        "Representative Multi-Channel Stage 2 Context Images",
        ["stage2_topo", "stage2_modulus", "stage2_adhesion", "stage2_deformation"],
        "Chapter6_Stage1_Results_Feasibility_DRAFT_v1alpha-20260306-figureimport.docx",
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
